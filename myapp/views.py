import base64
import binascii
import io
import json
import logging
import mimetypes
import os
import re
import secrets
import time
from pathlib import Path
from decimal import Decimal, InvalidOperation
from functools import wraps
from urllib.parse import unquote, urlencode, urlsplit

import requests
from docx import Document as WordDocument
from PIL import Image, ImageOps

from django.contrib import messages
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.models import User
from django.db import OperationalError, ProgrammingError
from django.db.models import Q, F, Count, Sum, Prefetch
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import resolve, Resolver404, reverse
from django.templatetags.static import static as static_url
from django.http import JsonResponse, StreamingHttpResponse, FileResponse, HttpResponse
from django.core.files.base import ContentFile
from django.core.cache import cache
from django.core.files.storage import default_storage
from django.core.management import call_command
from django.conf import settings
from django.utils import timezone
from datetime import timedelta, timezone as dt_timezone
from myapp.forms import (
    AISignupForm, PhoneVerifyForm, AILoginForm, SignupEditForm,
    AIProfileEditForm, AIPasswordChangeForm, CategoryForm, OrderStatusForm,
    ProductForm, ProductImageFormSet, ProductColorFormSet,
    AboutUsContentForm, PolicyPageForm, PaymentSettingsForm, DropboxSettingsForm, PWASettingsForm,
    FeeSettingsForm, GrantAISubscriptionForm, AddUserForm, SiteCustomizationForm,
)
from myapp.models import (
    ContactLead, StoreProfile, Category, Order, OrderItem,
    Product, AboutUsContent, PolicyPage, PaymentSettings, Payment,
    DropboxSettings, PhoneVerification, PWASettings, FeeSettings, SiteCustomization,
    AIConversation, AIMessage, AIBlock, AINote, AIReport, AIGeneratedFile,
    GitHubConnection, YouTubeDownloadJob,
)
from myapp import dropbox_backup
from myapp import ai_chat
from myapp import github_ops
from myapp import doc_extract
from myapp import company_knowledge
from myapp import image_ocr
from myapp import image_generation
from myapp import privacy
from myapp import request_router
from myapp import audio_transcribe
from myapp import youtube_download
from myapp.ai_report_analysis import analyze_report, aggregate_report_issues
from myapp.emailing import send_store_email, get_notify_email
from myapp.sms import send_phone_otp, verify_phone_otp
from myapp.seed_data import seed_demo_reviews
from myapp.single_device import register_active_session

logger = logging.getLogger(__name__)

try:
    import razorpay
except ImportError:  # pragma: no cover - optional dependency until configured
    razorpay = None


def site_customization_context(request):
    """Registered in edutrellis/settings.py TEMPLATES/OPTIONS/context_
    processors — exposes the admin-uploaded favicon (dashboard → Customize)
    to every template as SITE_FAVICON_URL, so each `<link rel="icon">` tag
    can fall back to the static default without every view needing to
    fetch SiteCustomization itself."""
    try:
        obj = SiteCustomization.get_solo()
        return {'SITE_FAVICON_URL': obj.favicon.url if obj.favicon else None}
    except (OperationalError, ProgrammingError):
        # A restored backup can predate the SiteCustomization migration.
        # Keep every page renderable until the post-restore migration step
        # below upgrades that older schema.
        return {'SITE_FAVICON_URL': None}


def pwa_service_worker(request):
    response = render(request, 'sw.js', content_type='application/javascript')
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Service-Worker-Allowed'] = '/'
    return response


def _user_payload(user):
    profile = getattr(user, 'store_profile', None)
    return {
        'name': user.get_full_name().strip() or user.username,
        'email': user.email,
        'phone': profile.phone if profile else '',
        'is_staff': user.is_staff,
        'is_superuser': user.is_superuser,
        'avatar_url': profile.avatar.url if (profile and profile.avatar) else None,
        'wallet_balance': float(profile.wallet_balance) if profile else 0.0,
        'phone_verified': bool(profile and profile.phone_verified),
        'location_consent': profile.location_consent if profile else StoreProfile.LOCATION_UNKNOWN,
        'amount_paid_missing': bool(not profile or profile.manual_amount_paid <= 0),
    }


def _location_prompt_needed(user):
    if not user.is_authenticated:
        return False
    profile = getattr(user, 'store_profile', None)
    return not profile or profile.location_consent == StoreProfile.LOCATION_UNKNOWN


def _profile_wizard_needed(user):
    """Whether to show the "some info needs to be updated" wizard (see
    includes/profile_wizard.html) — an older account still missing BOTH
    its location and a recorded amount-paid figure. Only fires when both
    are missing since the wizard's fixed two-step flow always asks for
    both; a user missing just one of the two (e.g. staff already recorded
    their amount paid) isn't shown it — the plain one-time location card
    (_location_prompt_needed) still covers a location-only gap on its own."""
    if not user.is_authenticated:
        return False
    profile = getattr(user, 'store_profile', None)
    location_missing = not profile or profile.location_consent == StoreProfile.LOCATION_UNKNOWN
    amount_missing = not profile or profile.manual_amount_paid <= 0
    return location_missing and amount_missing


def _merge_session_ai_chats_into_user(user, session_key):
    """Attach a guest's saved AI data to the account they just entered."""
    if not session_key:
        return
    AIConversation.objects.filter(session_key=session_key, user__isnull=True).update(user=user, session_key='')
    AIGeneratedFile.objects.filter(session_key=session_key, user__isnull=True).update(user=user, session_key='')
    # A report submitted before login should not disappear from the account
    # panel afterwards.  Keep its immutable evidence and status, but move its
    # owner from the guest session to the authenticated account just like the
    # conversation itself.
    AIReport.objects.filter(session_key=session_key, user__isnull=True).update(user=user, session_key='')


def _parse_json_body(request):
    try:
        return json.loads(request.body or '{}')
    except json.JSONDecodeError:
        return {}


PHONE_VERIFY_OTP_TTL_MINUTES = 10
PHONE_VERIFY_RESEND_COOLDOWN_SECONDS = 45
PHONE_VERIFY_MAX_ATTEMPTS = 5


def _new_phone_verification(user, phone, now):
    """Sends a fresh OTP via 2Factor and records the resulting session so
    a later confirm can check the code against it. The OTP digits
    themselves live at 2Factor, not here."""
    session_id = send_phone_otp(phone)
    pending, _created = PhoneVerification.objects.update_or_create(
        user=user,
        defaults={
            'session_id': session_id, 'phone': phone, 'attempts': 0, 'last_sent_at': now,
            'expires_at': now + timedelta(minutes=PHONE_VERIFY_OTP_TTL_MINUTES),
        },
    )
    if settings.DEBUG:
        # Local/dev visibility — the OTP itself isn't known to us (2Factor
        # generates and checks it), but this confirms the SMS send attempt
        # ran and which phone/session it's tied to.
        print(f"\n{'='*60}\nPhone OTP sent to {phone} (2Factor session {session_id})\n{'='*60}\n")
    return pending


def ai_signup(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    form = AISignupForm(_parse_json_body(request))
    if not form.is_valid():
        return JsonResponse(
            {'status': 'validation_error', 'errors': {k: v[0] for k, v in form.errors.items()}},
            status=400,
        )

    name = form.cleaned_data['name']
    phone = form.cleaned_data['phone']
    email = form.cleaned_data['email']
    password = form.cleaned_data['password']
    first_name, _, last_name = name.partition(' ')

    user = User.objects.create_user(
        username=email, email=email, password=password,
        first_name=first_name, last_name=last_name,
    )
    StoreProfile.objects.create(user=user, phone=phone)

    if not request.session.session_key:
        request.session.create()
    pre_login_session_key = request.session.session_key

    auth_user = authenticate(request, username=email, password=password)
    if auth_user:
        login(request, auth_user)
        _merge_session_ai_chats_into_user(auth_user, pre_login_session_key)

    # Best-effort — the account is already created and the user is
    # already logged in above, so a slow/flaky SMS send (or none configured
    # at all) can never fail or delay signup itself. They can verify anytime
    # from Edit Profile.
    try:
        _new_phone_verification(auth_user or user, phone, timezone.now())
        print(f"[signup sms] OTP SMS SENT via 2Factor to {phone}")
    except Exception as e:
        import traceback
        print(f"[signup sms] OTP SMS FAILED to send via 2Factor to {phone}: {e!r}")
        traceback.print_exc()
        logger.warning("Signup verification SMS failed for %s: %s", phone, e)

    return JsonResponse({'status': 'ok', 'user': _user_payload(auth_user or user)})


def ai_phone_verify_send(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    if profile.phone_verified:
        return JsonResponse({'status': 'error', 'detail': 'Your phone number is already verified.'}, status=400)
    if not profile.phone:
        return JsonResponse({'status': 'error', 'detail': 'Add a phone number to your profile first.'}, status=400)

    now = timezone.now()
    existing = PhoneVerification.objects.filter(user=request.user).first()
    if existing and (now - existing.last_sent_at).total_seconds() < PHONE_VERIFY_RESEND_COOLDOWN_SECONDS:
        wait = PHONE_VERIFY_RESEND_COOLDOWN_SECONDS - int((now - existing.last_sent_at).total_seconds())
        return JsonResponse({'status': 'error', 'detail': f'Please wait {wait}s before requesting another code.'}, status=429)

    try:
        _new_phone_verification(request.user, profile.phone, now)
    except Exception as e:
        logger.exception("Phone verification SMS failed for %s: %s", profile.phone, e)
        return JsonResponse({'status': 'error', 'detail': 'Could not send the verification code. Please try again shortly.'}, status=502)

    return JsonResponse({'status': 'ok'})


def ai_phone_verify_confirm(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    form = PhoneVerifyForm(_parse_json_body(request))
    if not form.is_valid():
        return JsonResponse({'status': 'validation_error', 'errors': {k: v[0] for k, v in form.errors.items()}}, status=400)

    otp = form.cleaned_data['otp']
    pending = PhoneVerification.objects.filter(user=request.user).first()
    if not pending:
        return JsonResponse({'status': 'error', 'detail': 'No pending verification — send a new code first.'}, status=404)

    if pending.is_expired:
        pending.delete()
        return JsonResponse({'status': 'error', 'detail': 'That code has expired — send a new one.'}, status=400)

    if pending.attempts >= PHONE_VERIFY_MAX_ATTEMPTS:
        pending.delete()
        return JsonResponse({'status': 'error', 'detail': 'Too many incorrect attempts — send a new code.'}, status=400)

    try:
        matched = verify_phone_otp(pending.session_id, otp)
    except Exception as e:
        logger.exception("2Factor verify call failed for %s: %s", pending.phone, e)
        return JsonResponse({'status': 'error', 'detail': 'Could not verify that code right now. Please try again shortly.'}, status=502)

    if not matched:
        pending.attempts += 1
        pending.save(update_fields=['attempts'])
        left = PHONE_VERIFY_MAX_ATTEMPTS - pending.attempts
        return JsonResponse({'status': 'error', 'detail': f'Incorrect code — {left} attempt{"s" if left != 1 else ""} left.'}, status=400)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    profile.phone_verified = True
    profile.save(update_fields=['phone_verified'])
    pending.delete()

    return JsonResponse({'status': 'ok', 'user': _user_payload(request.user)})


def ai_login(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    form = AILoginForm(_parse_json_body(request))
    if not form.is_valid():
        return JsonResponse(
            {'status': 'validation_error', 'errors': {k: v[0] for k, v in form.errors.items()}},
            status=400,
        )

    identifier = form.cleaned_data['identifier'].strip()
    password = form.cleaned_data['password']

    user_obj = User.objects.filter(email__iexact=identifier).first()
    if not user_obj:
        profile = StoreProfile.objects.filter(phone=identifier).first()
        user_obj = profile.user if profile else None

    if not user_obj:
        return JsonResponse({'status': 'error', 'detail': 'No account found with that email or phone.'}, status=400)

    auth_user = authenticate(request, username=user_obj.username, password=password)
    if not auth_user:
        return JsonResponse({'status': 'error', 'detail': 'Incorrect password.'}, status=400)

    if not request.session.session_key:
        request.session.create()
    pre_login_session_key = request.session.session_key

    login(request, auth_user)
    _merge_session_ai_chats_into_user(auth_user, pre_login_session_key)

    return JsonResponse({'status': 'ok', 'user': _user_payload(auth_user)})


def ai_location_update(request):
    """Save one consent choice and, when allowed, one browser location fix."""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    data = _parse_json_body(request)
    consent = data.get('consent')
    if consent not in (StoreProfile.LOCATION_GRANTED, StoreProfile.LOCATION_DENIED):
        return JsonResponse({'status': 'error', 'detail': 'Choose whether to enable location.'}, status=400)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    profile.location_consent = consent
    profile.location_updated_at = timezone.now()
    update_fields = [
        'location_consent', 'location_updated_at', 'location_latitude',
        'location_longitude', 'location_accuracy_m',
    ]

    if consent == StoreProfile.LOCATION_DENIED:
        profile.location_latitude = None
        profile.location_longitude = None
        profile.location_accuracy_m = None
    else:
        try:
            latitude = Decimal(str(data.get('latitude')))
            longitude = Decimal(str(data.get('longitude')))
            accuracy = Decimal(str(data.get('accuracy', 0)))
            if not latitude.is_finite() or not longitude.is_finite() or not accuracy.is_finite():
                raise InvalidOperation
            if not Decimal('-90') <= latitude <= Decimal('90'):
                raise InvalidOperation
            if not Decimal('-180') <= longitude <= Decimal('180'):
                raise InvalidOperation
            if not Decimal('0') <= accuracy <= Decimal('1000000'):
                raise InvalidOperation
        except (InvalidOperation, TypeError, ValueError):
            return JsonResponse({'status': 'error', 'detail': 'The browser returned an invalid location.'}, status=400)
        profile.location_latitude = latitude
        profile.location_longitude = longitude
        profile.location_accuracy_m = int(accuracy.to_integral_value())

    profile.save(update_fields=update_fields)
    return JsonResponse({
        'status': 'ok',
        'location_consent': profile.location_consent,
        'location_captured': bool(
            profile.location_latitude is not None and profile.location_longitude is not None
        ),
    })


AMOUNT_PAID_MAX = Decimal('10000000')


def ai_amount_paid_update(request):
    """Lets a logged-in user self-report how much they've already paid, for
    older accounts staff hasn't recorded a manual_amount_paid figure for
    yet (see the profile-completion wizard, includes/profile_wizard.html).
    Only accepted while the field is still at its untouched default of 0 —
    once a real value is on file (self-reported here, or entered by staff
    in the dashboard), further self-reports are silently ignored so this
    can't be used to overwrite that figure after the fact from the client."""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    data = _parse_json_body(request)
    try:
        amount = Decimal(str(data.get('amount')))
        if not amount.is_finite() or amount < 0 or amount > AMOUNT_PAID_MAX:
            raise InvalidOperation
    except (InvalidOperation, TypeError, ValueError):
        return JsonResponse({'status': 'error', 'detail': 'Enter a valid amount.'}, status=400)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    if profile.manual_amount_paid <= 0:
        profile.manual_amount_paid = amount
        profile.save(update_fields=['manual_amount_paid'])
    return JsonResponse({'status': 'ok', 'manual_amount_paid': float(profile.manual_amount_paid)})


def ai_logout(request):
    logout(request)
    return JsonResponse({'status': 'ok'})


def ai_profile_update(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    form = AIProfileEditForm(request.POST, request.FILES, user=request.user)
    if not form.is_valid():
        return JsonResponse(
            {'status': 'validation_error', 'errors': {k: v[0] for k, v in form.errors.items()}},
            status=400,
        )

    name = form.cleaned_data['name']
    first_name, _, last_name = name.partition(' ')
    request.user.first_name = first_name
    request.user.last_name = last_name
    user_fields = ['first_name', 'last_name']
    email = form.cleaned_data.get('email')
    if email:
        # Login resolves an account by its current email and then authenticates
        # its internal username, so staff/custom usernames do not need changing.
        request.user.email = email
        user_fields.append('email')
    request.user.save(update_fields=user_fields)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    new_phone = form.cleaned_data['phone']
    if profile.phone != new_phone:
        profile.phone_verified = False
    profile.phone = new_phone
    if form.cleaned_data.get('avatar'):
        profile.avatar = form.cleaned_data['avatar']
    profile.save()

    return JsonResponse({'status': 'ok', 'user': _user_payload(request.user)})


def ai_password_change(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    form = AIPasswordChangeForm(_parse_json_body(request))
    if not form.is_valid():
        return JsonResponse(
            {'status': 'validation_error', 'errors': {k: v[0] for k, v in form.errors.items()}},
            status=400,
        )

    if not request.user.check_password(form.cleaned_data['current_password']):
        return JsonResponse({'status': 'error', 'detail': 'Current password is incorrect.'}, status=400)

    request.user.set_password(form.cleaned_data['new_password'])
    request.user.save(update_fields=['password'])
    update_session_auth_hash(request, request.user)  # keep the session logged in
    register_active_session(request.user, request.session.session_key)
    return JsonResponse({'status': 'ok'})


def ai_account_details(request):
    """Return the signed-in customer's account data for the AI profile menu.

    Reports are intentionally selected through request.user only; guest
    reports and reports belonging to other accounts must never appear here.
    """
    if request.method != 'GET':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'detail': 'You need to be logged in.'}, status=401)

    profile, _ = StoreProfile.objects.get_or_create(user=request.user)
    is_staff = request.user.is_staff
    is_subscribed = profile.is_ai_subscribed
    if is_staff:
        plan_name = 'Staff access'
    elif is_subscribed:
        plan_name = 'Vidhyora AI Premium'
    else:
        plan_name = 'Free plan'

    reports = []
    for report in AIReport.objects.filter(user=request.user).select_related('conversation')[:100]:
        reports.append({
            'id': report.pk,
            'status': report.status,
            'status_label': report.get_status_display(),
            'explanation': report.explanation,
            'user_prompt': report.user_prompt,
            'reported_reply': report.reported_reply,
            'model': ai_chat.MODELS.get(report.model_key, {}).get('label', report.model_key or 'AI'),
            'conversation_title': report.conversation.title if report.conversation else 'Deleted conversation',
            'created_at': timezone.localtime(report.created_at).isoformat(),
        })

    # Every FLUX-generated/edited image across all of this user's own
    # conversations, newest first — the "My Images" gallery. image_data on
    # an assistant turn is only ever the real stored URL (see
    # _ai_flux_response), never trusted/AI-invented text. Do not filter by
    # model_key here: when ChatGPT 5.6 initiated the generation, that public
    # identity is stored instead of exposing the internal FLUX worker.
    images = [
        {
            'id': message.pk,
            'url': message.image_data,
            'conversation_id': message.conversation_id,
            'created_at': timezone.localtime(message.created_at).isoformat(),
        }
        for message in AIMessage.objects.filter(
            conversation__user=request.user,
            role=AIMessage.ROLE_ASSISTANT,
        ).exclude(image_data='').order_by('-created_at')[:60]
    ]

    response = JsonResponse({
        'status': 'ok',
        'user': _user_payload(request.user),
        'subscription': {
            'plan_name': plan_name,
            'active': bool(is_staff or is_subscribed),
            'is_staff': is_staff,
            'expires_at': (
                timezone.localtime(profile.ai_subscription_until).isoformat()
                if profile.ai_subscription_until and is_subscribed else None
            ),
            'free_used': profile.ai_free_messages_used,
            'free_limit': AI_FREE_MESSAGE_LIMIT,
            'free_remaining': max(0, AI_FREE_MESSAGE_LIMIT - profile.ai_free_messages_used),
            'purchase_url': _ai_purchase_url(),
        },
        'reports': reports,
        'images': images,
    })
    response['Cache-Control'] = 'private, no-store'
    return response


def custom_404(request, exception=None):
    # edutrellis/urls.py adds a catch-all pattern (matching every path, with
    # or without a trailing slash) so this view renders even with DEBUG=True.
    # That catch-all is itself a URL match, though, so it silently defeats
    # Django's normal APPEND_SLASH redirect (which only fires when the
    # un-slashed path doesn't resolve to anything). Reimplement that check
    # here against myapp.urls directly (which has no catch-all).
    path = request.path
    # Only redirect safe methods — a 302 on POST/PUT/etc. gets turned into a
    # GET by the browser, silently dropping the request body.
    if request.method in ('GET', 'HEAD') and not path.endswith('/'):
        try:
            resolve(path + '/', urlconf='myapp.urls')
            query = f'?{request.META["QUERY_STRING"]}' if request.META.get('QUERY_STRING') else ''
            return redirect(path + '/' + query, permanent=True)
        except Resolver404:
            pass
    return render(request, '404.html', status=404)


def _dashboard_guard(request):
    """Only authenticated staff can use the custom dashboard."""
    return request.user.is_authenticated and request.user.is_staff


def dashboard_staff_required(view_func):
    """Every dashboard_* view needs this same staff-only check — used to be
    copy-pasted as the first two lines of each one (easy to forget on a new
    view, silently exposing a staff-only page). Applying this decorator
    instead makes the guard structurally impossible to skip."""
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not _dashboard_guard(request):
            return redirect('ai_page')
        return view_func(request, *args, **kwargs)
    return wrapper


@dashboard_staff_required
def dashboard_home(request):
    now = timezone.now()
    conversations = AIConversation.objects.select_related('user').annotate(message_count=Count('messages'))
    reports = AIReport.objects.select_related('user', 'conversation')
    total_conversations = conversations.count()
    total_messages = AIMessage.objects.count()
    registered_conversations = conversations.exclude(user__isnull=True).count()
    guest_conversations = conversations.filter(user__isnull=True).count()
    open_reports = reports.filter(status=AIReport.STATUS_OPEN).count()
    resolved_reports = reports.filter(status=AIReport.STATUS_RESOLVED).count()
    total_accounts = User.objects.count()
    location_enabled = StoreProfile.objects.filter(location_consent=StoreProfile.LOCATION_GRANTED).count()
    location_declined = StoreProfile.objects.filter(location_consent=StoreProfile.LOCATION_DENIED).count()
    location_not_asked = max(0, total_accounts - location_enabled - location_declined)
    location_enabled_pct = round((location_enabled / total_accounts) * 100) if total_accounts else 0
    location_declined_end_pct = min(
        100,
        location_enabled_pct + (round((location_declined / total_accounts) * 100) if total_accounts else 0),
    )

    today = timezone.localdate()
    chart_start = today - timedelta(days=6)

    def daily_counts(queryset, field):
        filters = {
            f'{field}__date__gte': chart_start,
            f'{field}__date__lte': today,
        }
        return dict(
            queryset.filter(**filters)
            .values_list(f'{field}__date').annotate(total=Count('id'))
        )

    conversation_counts = daily_counts(AIConversation.objects.all(), 'created_at')
    message_counts = daily_counts(AIMessage.objects.all(), 'created_at')
    activity_max = max([1, *conversation_counts.values(), *message_counts.values()])
    daily_activity = []
    for offset in range(7):
        day = chart_start + timedelta(days=offset)
        conversation_count = conversation_counts.get(day, 0)
        message_count = message_counts.get(day, 0)
        daily_activity.append({
            'label': day.strftime('%d %b'),
            'conversations': conversation_count,
            'messages': message_count,
            'conversation_height': round((conversation_count / activity_max) * 100),
            'message_height': round((message_count / activity_max) * 100),
        })

    report_total = open_reports + resolved_reports
    report_open_pct = round((open_reports / report_total) * 100) if report_total else 0
    registered_pct = round((registered_conversations / total_conversations) * 100) if total_conversations else 0
    guest_pct = 100 - registered_pct if total_conversations else 0

    model_rows = list(
        AIMessage.objects.filter(role=AIMessage.ROLE_ASSISTANT).exclude(model_key='')
        .values('model_key').annotate(total=Count('id')).order_by('-total', 'model_key')[:5]
    )
    model_max = max((row['total'] for row in model_rows), default=0)
    model_usage = [
        {
            'key': row['model_key'],
            'label': ai_chat.MODELS.get(row['model_key'], {}).get('label', row['model_key']),
            'total': row['total'],
            'width': round((row['total'] / model_max) * 100) if model_max else 0,
        }
        for row in model_rows
    ]
    context = {
        'active': 'home',
        'total_conversations': total_conversations,
        'total_messages': total_messages,
        'registered_ai_users': conversations.exclude(user__isnull=True).values('user_id').distinct().count(),
        'guest_conversations': guest_conversations,
        'active_subscribers': StoreProfile.objects.filter(ai_subscription_until__gt=now).count(),
        'open_reports': open_reports,
        'resolved_reports': resolved_reports,
        'active_blocks': AIBlock.objects.count(),
        'daily_activity': daily_activity,
        'registered_conversations': registered_conversations,
        'registered_pct': registered_pct,
        'guest_pct': guest_pct,
        'report_total': report_total,
        'report_open_pct': report_open_pct,
        'model_usage': model_usage,
        'location_accounts': total_accounts,
        'location_enabled': location_enabled,
        'location_declined': location_declined,
        'location_not_asked': location_not_asked,
        'location_enabled_pct': location_enabled_pct,
        'location_declined_end_pct': location_declined_end_pct,
        'recent_conversations': conversations.order_by('-updated_at')[:8],
        'recent_reports': reports.order_by('-created_at')[:8],
    }
    return render(request, 'dashboard/home.html', context)


@dashboard_staff_required
def dashboard_signups(request):
    q = request.GET.get('q', '').strip()
    users = User.objects.select_related('store_profile').order_by('-date_joined')
    if q:
        users = users.filter(
            Q(username__icontains=q) | Q(email__icontains=q) |
            Q(first_name__icontains=q) | Q(last_name__icontains=q) |
            Q(store_profile__phone__icontains=q)
        )
    today = timezone.localdate()
    chart_start = today - timedelta(days=6)
    counts_by_day = dict(
        User.objects.filter(date_joined__date__gte=chart_start, date_joined__date__lte=today)
        .values_list('date_joined__date').annotate(total=Count('id'))
    )
    chart_counts = [counts_by_day.get(chart_start + timedelta(days=offset), 0) for offset in range(7)]
    chart_max = max(chart_counts) if chart_counts else 0
    signup_chart = [
        {
            'label': (chart_start + timedelta(days=offset)).strftime('%d %b'),
            'count': count,
            'height': round((count / chart_max) * 100) if chart_max else 0,
        }
        for offset, count in enumerate(chart_counts)
    ]
    total_signups = User.objects.count()
    signups_today = User.objects.filter(date_joined__date=today).count()
    signups_last_7_days = User.objects.filter(date_joined__date__gte=chart_start).count()
    signups_this_month = User.objects.filter(
        date_joined__date__gte=today.replace(day=1), date_joined__date__lte=today,
    ).count()
    total_amount_paid = StoreProfile.objects.aggregate(total=Sum('manual_amount_paid'))['total'] or 0
    paid_users = StoreProfile.objects.filter(manual_amount_paid__gt=0).count()
    no_recorded_payment = max(0, total_signups - paid_users)
    paid_users_pct = round((paid_users / total_signups) * 100) if total_signups else 0

    superuser_count = User.objects.filter(is_superuser=True).count()
    staff_count = User.objects.filter(is_staff=True, is_superuser=False).count()
    customer_count = User.objects.filter(is_staff=False, is_superuser=False).count()
    role_max = max(1, customer_count, staff_count, superuser_count)
    account_roles = [
        {'label': 'Customers', 'count': customer_count, 'width': round((customer_count / role_max) * 100)},
        {'label': 'Staff', 'count': staff_count, 'width': round((staff_count / role_max) * 100)},
        {'label': 'Superusers', 'count': superuser_count, 'width': round((superuser_count / role_max) * 100)},
    ]
    return render(request, 'dashboard/signups.html', {
        'active': 'signups', 'users': users, 'q': q, 'add_user_form': AddUserForm(),
        'total_signups': total_signups, 'total_amount_paid': total_amount_paid,
        'signups_today': signups_today, 'signups_last_7_days': signups_last_7_days,
        'signups_this_month': signups_this_month, 'signup_chart': signup_chart,
        'paid_users': paid_users, 'no_recorded_payment': no_recorded_payment,
        'paid_users_pct': paid_users_pct, 'account_roles': account_roles,
        'location_enabled': StoreProfile.objects.filter(
            location_consent=StoreProfile.LOCATION_GRANTED,
        ).count(),
    })


@dashboard_staff_required
def dashboard_user_add(request):
    """Manually create a customer account from the dashboard — used by both
    the Signups page and AI Management (so staff can create someone to
    grant AI premium to without them self-registering first)."""
    next_url = request.POST.get('next', '')
    if next_url not in ('dashboard_signups', 'dashboard_ai_management'):
        next_url = 'dashboard_signups'
    if request.method == 'POST':
        form = AddUserForm(request.POST)
        if form.is_valid():
            name = form.cleaned_data['name']
            email = form.cleaned_data['email']
            phone = form.cleaned_data['phone']
            amount_paid = form.cleaned_data['amount_paid']
            password = form.cleaned_data['password'] or secrets.token_urlsafe(9)
            first_name, _, last_name = name.partition(' ')
            user = User.objects.create_user(
                username=email, email=email, password=password,
                first_name=first_name, last_name=last_name,
            )
            StoreProfile.objects.create(user=user, phone=phone, manual_amount_paid=amount_paid)
            if form.cleaned_data['password']:
                messages.success(request, f'Created account for {email}.')
            else:
                messages.success(request, f'Created account for {email} — temporary password: {password}')
        else:
            for errs in form.errors.values():
                for error in errs:
                    messages.error(request, error)
    return redirect(next_url)


@dashboard_staff_required
def dashboard_signup_edit(request, pk):
    edited_user = get_object_or_404(User, pk=pk)
    profile, _ = StoreProfile.objects.get_or_create(user=edited_user)
    form = SignupEditForm(
        request.POST or None, instance=edited_user,
        initial={
            'phone': profile.phone, 'wallet_balance': profile.wallet_balance,
            'amount_paid': profile.manual_amount_paid,
        },
    )
    if request.method == 'POST' and form.is_valid():
        form.save()
        profile.phone = form.cleaned_data['phone']
        profile.wallet_balance = form.cleaned_data['wallet_balance']
        profile.manual_amount_paid = form.cleaned_data['amount_paid']
        profile.save(update_fields=['phone', 'wallet_balance', 'manual_amount_paid'])
        return redirect('dashboard_signups')
    return render(request, 'dashboard/signup_form.html', {'active': 'signups', 'form': form, 'edited_user': edited_user})


@dashboard_staff_required
def dashboard_signup_delete(request, pk):
    if request.method == 'POST':
        target = get_object_or_404(User, pk=pk)
        if target.pk == request.user.pk:
            messages.error(request, "You can't delete your own account from here.")
        elif target.is_staff or target.is_superuser:
            messages.error(request, "Staff and admin accounts can't be deleted from here — use Django admin if you're sure.")
        elif Order.objects.filter(user=target).exists():
            # Deleting the User cascades and wipes their Order/OrderItem/
            # Payment rows — real financial records, not just a login.
            messages.error(request, "This customer has order history — deleting the account would erase those orders and payment records. Use Django admin if you're sure.")
        else:
            target.delete()
            messages.success(request, 'Customer account deleted.')
    return redirect('dashboard_signups')


@dashboard_staff_required
def dashboard_ai_management(request):
    """Lets staff see who has Vidhyora AI access and for how long, and
    manually grant a customer premium access (see GrantAISubscriptionForm)
    without them having to actually buy the plan — e.g. a comped account.
    Deliberately separate from the Django-admin/dashboard-staff tier: this
    only ever touches StoreProfile.ai_subscription_until, never
    is_staff/is_superuser, so granting someone AI premium here can never
    accidentally hand them dashboard or Django-admin access."""
    now = timezone.now()
    subscribers = (
        StoreProfile.objects.filter(ai_subscription_until__isnull=False)
        .select_related('user').order_by('-ai_subscription_until')
    )
    admins = User.objects.filter(Q(is_staff=True) | Q(is_superuser=True)).select_related('store_profile').order_by('-date_joined')
    context = {
        'active': 'ai_management',
        'subscribers': subscribers,
        'active_subscriber_count': subscribers.filter(ai_subscription_until__gt=now).count(),
        'admins': admins,
        'now': now,
        'grant_form': GrantAISubscriptionForm(),
        'add_user_form': AddUserForm(),
    }
    return render(request, 'dashboard/ai_management.html', context)


@dashboard_staff_required
def dashboard_ai_grant(request):
    if request.method == 'POST':
        form = GrantAISubscriptionForm(request.POST)
        if form.is_valid():
            target_user = form.matched_user
            days = form.cleaned_data['days']
            profile, _ = StoreProfile.objects.get_or_create(user=target_user)
            profile.ai_subscription_until = timezone.now() + timedelta(days=days)
            profile.ai_free_messages_used = 0
            profile.save(update_fields=['ai_subscription_until', 'ai_free_messages_used'])
            messages.success(
                request,
                f"Granted {target_user.email or target_user.username} Vidhyora AI premium access "
                f"until {timezone.localtime(profile.ai_subscription_until):%d %b %Y}.",
            )
        else:
            for error in form.errors.get('identifier', []):
                messages.error(request, error)
            for error in form.errors.get('days', []):
                messages.error(request, error)
    return redirect('dashboard_ai_management')


@dashboard_staff_required
def dashboard_ai_revoke(request, pk):
    if request.method == 'POST':
        profile = get_object_or_404(StoreProfile, pk=pk)
        profile.ai_subscription_until = None
        profile.save(update_fields=['ai_subscription_until'])
        messages.success(request, f'Revoked Vidhyora AI premium access for {profile.user.email or profile.user.username}.')
    return redirect('dashboard_ai_management')


def _ai_activity_redirect(conversation_id):
    if conversation_id:
        try:
            return redirect('dashboard_ai_activity_detail', pk=int(conversation_id))
        except (TypeError, ValueError):
            pass
    return redirect('dashboard_ai_activity')


@dashboard_staff_required
def dashboard_ai_activity(request):
    """Every AI conversation — who sent it (account or guest IP) and how
    many messages — so staff can actually see a spam pattern (same IP or
    account hammering the chat) instead of it sitting invisibly behind the
    live rate limiter. See AIBlock / dashboard_ai_block for the
    accompanying block tools."""
    q = request.GET.get('q', '').strip()
    conversations = (
        AIConversation.objects.select_related('user')
        .annotate(message_count=Count('messages'))
        .order_by('-updated_at')
    )
    if q:
        conversations = conversations.filter(
            Q(user__email__icontains=q) | Q(user__username__icontains=q) |
            Q(ip_address__icontains=q) | Q(title__icontains=q)
        )
    blocked_ips = set(AIBlock.objects.exclude(ip_address__isnull=True).values_list('ip_address', flat=True))
    blocked_user_ids = set(AIBlock.objects.exclude(user__isnull=True).values_list('user_id', flat=True))
    context = {
        'active': 'ai_activity',
        'conversations': conversations[:200],
        'q': q,
        'blocks': AIBlock.objects.select_related('user', 'created_by').order_by('-created_at'),
        'blocked_ips': blocked_ips,
        'blocked_user_ids': blocked_user_ids,
    }
    return render(request, 'dashboard/ai_activity.html', context)


@dashboard_staff_required
def dashboard_ai_activity_detail(request, pk):
    conversation = get_object_or_404(AIConversation.objects.select_related('user'), pk=pk)
    context = {
        'active': 'ai_activity',
        'conversation': conversation,
        'ai_messages': conversation.messages.order_by('created_at'),
        'is_ip_blocked': bool(conversation.ip_address) and AIBlock.objects.filter(ip_address=conversation.ip_address).exists(),
        'is_user_blocked': bool(conversation.user_id) and AIBlock.objects.filter(user_id=conversation.user_id).exists(),
    }
    return render(request, 'dashboard/ai_activity_detail.html', context)


def _ai_reporter_scope(report):
    """Return the exact account/session identity represented by a report."""
    conversation = report.conversation
    user = report.user or (conversation.user if conversation and conversation.user_id else None)
    session_key = report.session_key or (
        conversation.session_key if conversation and not user else ''
    )
    return user, session_key


def _ai_report_system_health():
    """Scan objective storage/linkage gaps that report text alone cannot show."""
    signals = []
    generated_images = list(
        AIMessage.objects.filter(role=AIMessage.ROLE_ASSISTANT)
        .exclude(image_data='').values_list('image_data', flat=True)
    )
    missing_images = 0
    for image_value in generated_images:
        storage_name = _ai_storage_name_from_url(image_value)
        if not storage_name:
            continue
        try:
            if not default_storage.exists(storage_name):
                missing_images += 1
        except Exception:
            # Storage availability is unknown, not proof that the image is
            # missing. Do not turn a connectivity problem into a false alert.
            pass
    if missing_images:
        signals.append({
            'key': 'missing_generated_media',
            'label': 'Generated image files are missing',
            'severity': 'high',
            'count': missing_images,
            'evidence': (
                f'{missing_images} of {len(generated_images)} stored generated-image '
                'messages point to media files that are no longer available.'
            ),
            'fix': (
                'Move generated media to durable storage and include media in backup/restore. '
                'Keep the database snapshot added to new image reports as review evidence.'
            ),
        })

    unlinked_reports = AIReport.objects.filter(message__isnull=True).count()
    if unlinked_reports:
        signals.append({
            'key': 'unlinked_report_evidence',
            'label': 'Reports are missing a stable response link',
            'severity': 'medium',
            'count': unlinked_reports,
            'evidence': f'{unlinked_reports} reports are not linked to an assistant message.',
            'fix': (
                'Keep submitting the exact message id for saved replies and persist structured '
                'rows for provider errors so every future report has an authoritative target.'
            ),
        })

    return signals


@dashboard_staff_required
def dashboard_ai_reports(request):
    """Every 'report this reply' submission from the AI chat (see the
    Report button under each reply, and ai_report_submit) — who reported
    it (account, with the email/login to follow up with, or guest
    session), what reply they flagged, and why — grouped by open/resolved
    so staff can see what still needs attention."""
    q = request.GET.get('q', '').strip()
    reports = AIReport.objects.select_related('user', 'conversation', 'message').order_by('-created_at')
    if q:
        reports = reports.filter(
            Q(user__email__icontains=q) | Q(user__username__icontains=q) |
            Q(session_key__icontains=q) | Q(explanation__icontains=q) |
            Q(reported_reply__icontains=q) | Q(user_prompt__icontains=q) |
            Q(user_document_name__icontains=q) | Q(model_key__icontains=q) |
            Q(conversation__title__icontains=q)
        )
    all_reports = list(reports)
    for report in all_reports:
        report.detected_issues = analyze_report(report)
        report.primary_issue = report.detected_issues[0] if report.detected_issues else None
    groups = [
        {'status': value, 'label': label, 'reports': [r for r in all_reports if r.status == value]}
        for value, label in AIReport.STATUS_CHOICES
    ]
    return render(request, 'dashboard/ai_reports.html', {
        'active': 'ai_reports', 'reports': all_reports, 'groups': groups, 'q': q,
        'failure_insights': aggregate_report_issues(all_reports),
        'system_health': _ai_report_system_health(),
    })


@dashboard_staff_required
def dashboard_ai_report_detail(request, pk):
    report = get_object_or_404(
        AIReport.objects.select_related('user', 'conversation__user', 'message'), pk=pk,
    )
    owner_user, owner_session = _ai_reporter_scope(report)

    if owner_user:
        conversation_filter = Q(user=owner_user)
        report_filter = Q(user=owner_user) | Q(user__isnull=True, conversation__user=owner_user)
        reporter_label = owner_user.get_full_name() or owner_user.email or owner_user.username
    elif owner_session:
        conversation_filter = Q(user__isnull=True, session_key=owner_session)
        report_filter = Q(user__isnull=True, session_key=owner_session)
        reporter_label = f'Guest session {owner_session[:10]}…'
    elif report.conversation_id:
        conversation_filter = Q(pk=report.conversation_id)
        report_filter = Q(pk=report.pk)
        reporter_label = 'Guest (identity unavailable)'
    else:
        conversation_filter = Q(pk__in=[])
        report_filter = Q(pk=report.pk)
        reporter_label = 'Deleted/unknown reporter'

    message_queryset = AIMessage.objects.only(
        'id', 'conversation_id', 'role', 'content', 'image_data',
        'document_name', 'document_text', 'model_key', 'created_at',
    ).order_by('created_at', 'pk')
    conversations = list(
        AIConversation.objects.filter(conversation_filter)
        .prefetch_related(Prefetch('messages', queryset=message_queryset, to_attr='report_messages'))
        .order_by('-updated_at', '-pk')
    )
    reporter_reports = list(
        AIReport.objects.filter(report_filter)
        .select_related('conversation', 'message')
        .order_by('-created_at', '-pk')
    )
    for reporter_report in reporter_reports:
        reporter_report.detected_issues = analyze_report(reporter_report)

    return render(request, 'dashboard/ai_report_detail.html', {
        'active': 'ai_reports',
        'report': report,
        'reporter_label': reporter_label,
        'current_issues': analyze_report(report),
        'reporter_insights': aggregate_report_issues(reporter_reports),
        'reporter_reports': reporter_reports,
        'conversations': conversations,
        'conversation_count': len(conversations),
        'message_count': sum(len(conversation.report_messages) for conversation in conversations),
    })


@dashboard_staff_required
def dashboard_ai_report_status_update(request, pk):
    if request.method == 'POST':
        report = get_object_or_404(AIReport, pk=pk)
        report.status = AIReport.STATUS_RESOLVED if report.status == AIReport.STATUS_OPEN else AIReport.STATUS_OPEN
        report.save(update_fields=['status'])
    return redirect('dashboard_ai_reports')


@dashboard_staff_required
def dashboard_ai_report_delete(request, pk):
    if request.method == 'POST':
        get_object_or_404(AIReport, pk=pk).delete()
    return redirect('dashboard_ai_reports')


def _dashboard_ai_image_response(image_value):
    """Serve private chat/report images without embedding base64 in HTML."""
    if not image_value:
        return HttpResponse('Image not available.', status=404, content_type='text/plain')

    if image_value.startswith('data:image/'):
        match = re.match(r'^data:(image/(?:png|jpe?g|webp|gif));base64,(.+)$', image_value, re.I | re.S)
        if not match:
            return HttpResponse('Image not available.', status=404, content_type='text/plain')
        try:
            raw = base64.b64decode(match.group(2), validate=True)
        except (binascii.Error, ValueError, TypeError):
            return HttpResponse('Image not available.', status=404, content_type='text/plain')
        if len(raw) > AI_IMAGE_MAX_DATA_URI_CHARS:
            return HttpResponse('Image is too large to preview.', status=413, content_type='text/plain')
        response = HttpResponse(raw, content_type=match.group(1).lower())
    else:
        storage_name = _ai_storage_name_from_url(image_value)
        if storage_name:
            try:
                image_file = default_storage.open(storage_name, 'rb')
            except Exception:
                return HttpResponse('Image file is unavailable.', status=404, content_type='text/plain')
            response = FileResponse(
                image_file,
                content_type=mimetypes.guess_type(storage_name)[0] or 'application/octet-stream',
            )
        else:
            parsed = urlsplit(image_value)
            if parsed.scheme not in ('http', 'https'):
                return HttpResponse('Image not available.', status=404, content_type='text/plain')
            return redirect(image_value)

    response['Cache-Control'] = 'private, no-store'
    response['X-Content-Type-Options'] = 'nosniff'
    return response


@dashboard_staff_required
def dashboard_ai_message_image(request, pk):
    message = get_object_or_404(AIMessage.objects.only('image_data'), pk=pk)
    return _dashboard_ai_image_response(message.image_data)


@dashboard_staff_required
def dashboard_ai_report_image(request, pk, side):
    report = get_object_or_404(AIReport.objects.only('user_image', 'reported_image'), pk=pk)
    if side == 'user':
        image_value = report.user_image
    elif side == 'assistant':
        image_value = report.reported_image
    else:
        return HttpResponse('Image not available.', status=404, content_type='text/plain')
    return _dashboard_ai_image_response(image_value)


@dashboard_staff_required
def dashboard_ai_block(request):
    if request.method == 'POST':
        ip_address = request.POST.get('ip_address', '').strip()
        user_id = request.POST.get('user_id', '').strip()
        conversation_id = request.POST.get('conversation_id', '').strip()
        reason = request.POST.get('reason', '').strip()
        target_user = None
        if user_id:
            target_user = get_object_or_404(User, pk=user_id)
            if target_user.is_staff or target_user.is_superuser:
                messages.error(request, "Staff/admin accounts can't be blocked from here.")
                return _ai_activity_redirect(conversation_id)
        if not ip_address and not target_user:
            messages.error(request, 'Nothing to block — no IP or account given.')
        else:
            block, created = AIBlock.objects.get_or_create(
                ip_address=ip_address or None, user=target_user,
                defaults={'reason': reason, 'created_by': request.user},
            )
            label = (target_user.email or target_user.username) if target_user else ip_address
            if created:
                messages.success(request, f'Blocked {label} from Vidhyora AI.')
            else:
                messages.error(request, f'{label} is already blocked.')
        return _ai_activity_redirect(conversation_id)
    return redirect('dashboard_ai_activity')


@dashboard_staff_required
def dashboard_ai_unblock(request, pk):
    if request.method == 'POST':
        block = get_object_or_404(AIBlock, pk=pk)
        label = (block.user.email or block.user.username) if block.user_id else block.ip_address
        block.delete()
        messages.success(request, f'Unblocked {label}.')
    return _ai_activity_redirect(request.POST.get('conversation_id', '').strip())


@dashboard_staff_required
def dashboard_contacts(request):
    q = request.GET.get('q', '').strip()
    leads = ContactLead.objects.order_by('-created_at')
    if q:
        leads = leads.filter(
            Q(name__icontains=q) | Q(email__icontains=q) | Q(phone__icontains=q) |
            Q(service__icontains=q) | Q(message__icontains=q)
        )

    all_leads = list(leads)
    groups = [
        {'source': value, 'label': label, 'leads': [l for l in all_leads if l.source == value]}
        for value, label in ContactLead.SOURCE_CHOICES
    ]
    return render(request, 'dashboard/contacts.html', {
        'active': 'contacts', 'leads': all_leads, 'groups': groups, 'q': q,
    })


@dashboard_staff_required
def dashboard_contact_delete(request, pk):
    if request.method == 'POST':
        get_object_or_404(ContactLead, pk=pk).delete()
    return redirect('dashboard_contacts')


@dashboard_staff_required
def dashboard_categories(request):
    q = request.GET.get('q', '').strip()
    categories = Category.objects.all()
    if q:
        categories = categories.filter(Q(name__icontains=q) | Q(slug__icontains=q))
    return render(request, 'dashboard/categories.html', {'active': 'categories', 'categories': categories, 'q': q})


@dashboard_staff_required
def dashboard_category_add(request):
    form = CategoryForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('dashboard_categories')
    return render(request, 'dashboard/category_form.html', {'active': 'categories', 'form': form, 'category': None})


@dashboard_staff_required
def dashboard_category_edit(request, pk):
    category = get_object_or_404(Category, pk=pk)
    form = CategoryForm(request.POST or None, request.FILES or None, instance=category)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('dashboard_categories')
    return render(request, 'dashboard/category_form.html', {'active': 'categories', 'form': form, 'category': category})


@dashboard_staff_required
def dashboard_category_delete(request, pk):
    if request.method == 'POST':
        get_object_or_404(Category, pk=pk).delete()
    return redirect('dashboard_categories')


@dashboard_staff_required
def dashboard_orders(request):
    q = request.GET.get('q', '').strip()
    status = request.GET.get('status', '').strip()
    orders = Order.objects.select_related('user').prefetch_related('items', 'payments').order_by('-created_at')
    if q:
        orders = orders.filter(
            Q(user__username__icontains=q) | Q(user__email__icontains=q) |
            Q(user__first_name__icontains=q) | Q(user__last_name__icontains=q)
        )
    if status:
        orders = orders.filter(status=status)
    return render(request, 'dashboard/orders.html', {
        'active': 'orders', 'orders': orders, 'q': q, 'status': status,
        'status_choices': Order.STATUS_CHOICES,
    })


@dashboard_staff_required
def dashboard_delivery(request):
    q = request.GET.get('q', '').strip()
    status = request.GET.get('status', '').strip()
    orders = Order.objects.select_related('user').prefetch_related('payments').filter(recipient_name__gt='').order_by('-created_at')
    if q:
        orders = orders.filter(
            Q(recipient_name__icontains=q) | Q(recipient_phone__icontains=q) |
            Q(city__icontains=q) | Q(pincode__icontains=q)
        )
    if status:
        orders = orders.filter(status=status)
    return render(request, 'dashboard/delivery.html', {
        'active': 'delivery', 'orders': orders, 'q': q, 'status': status,
        'status_choices': Order.STATUS_CHOICES,
    })


@dashboard_staff_required
def dashboard_order_status_update(request, pk):
    order = get_object_or_404(Order, pk=pk)
    if request.method == 'POST':
        form = OrderStatusForm(request.POST, instance=order)
        if form.is_valid():
            form.save()
            order.maybe_credit_wallet()
            order.maybe_grant_ai_subscription()
    return redirect('dashboard_orders')


@dashboard_staff_required
def dashboard_products(request):
    q = request.GET.get('q', '').strip()
    products = Product.objects.select_related('category').all()
    if q:
        products = products.filter(
            Q(name__icontains=q) | Q(brand__icontains=q) | Q(slug__icontains=q) | Q(tags__icontains=q)
        )
    return render(request, 'dashboard/products.html', {'active': 'products', 'products': products, 'q': q})


@dashboard_staff_required
def dashboard_product_add(request):
    form = ProductForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        product = form.save()
        # Images/video/colors are added on the edit page, once the product
        # (and therefore the FK the image/color formsets need) exists.
        return redirect('dashboard_product_edit', pk=product.pk)
    return render(request, 'dashboard/product_form.html', {
        'active': 'products', 'form': form, 'product': None,
        'image_formset': None, 'color_formset': None,
    })


@dashboard_staff_required
def dashboard_product_edit(request, pk):
    product = get_object_or_404(Product, pk=pk)
    form = ProductForm(request.POST or None, request.FILES or None, instance=product)
    image_formset = ProductImageFormSet(request.POST or None, request.FILES or None, instance=product, prefix='images')
    color_formset = ProductColorFormSet(request.POST or None, request.FILES or None, instance=product, prefix='colors')
    if request.method == 'POST' and form.is_valid() and image_formset.is_valid() and color_formset.is_valid():
        form.save()
        image_formset.save()
        color_formset.save()
        return redirect('dashboard_products')
    return render(request, 'dashboard/product_form.html', {
        'active': 'products', 'form': form, 'product': product,
        'image_formset': image_formset, 'color_formset': color_formset,
    })


@dashboard_staff_required
def dashboard_product_delete(request, pk):
    if request.method == 'POST':
        get_object_or_404(Product, pk=pk).delete()
    return redirect('dashboard_products')


@dashboard_staff_required
def dashboard_seed_reviews(request):
    if request.method == 'POST':
        messages.success(request, seed_demo_reviews())
    return redirect('dashboard_products')


@dashboard_staff_required
def dashboard_about(request):
    about = AboutUsContent.get_solo()
    form = AboutUsContentForm(request.POST or None, request.FILES or None, instance=about)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = AboutUsContentForm(instance=about)
    return render(request, 'dashboard/about_form.html', {'active': 'about', 'form': form, 'about': about, 'saved': saved})


@dashboard_staff_required
def dashboard_policies(request):
    policies = PolicyPage.objects.all()
    return render(request, 'dashboard/policies.html', {'active': 'policies', 'policies': policies})


@dashboard_staff_required
def dashboard_policy_edit(request, pk):
    policy = get_object_or_404(PolicyPage, pk=pk)
    form = PolicyPageForm(request.POST or None, instance=policy)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('dashboard_policies')
    return render(request, 'dashboard/policy_form.html', {'active': 'policies', 'form': form, 'policy': policy})


@dashboard_staff_required
def dashboard_payment_settings(request):
    settings_obj = PaymentSettings.get_solo()
    form = PaymentSettingsForm(request.POST or None, instance=settings_obj)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = PaymentSettingsForm(instance=settings_obj)
    return render(request, 'dashboard/payment_settings.html', {
        'active': 'payment_settings', 'form': form, 'settings_obj': settings_obj,
        'razorpay_installed': razorpay is not None, 'saved': saved,
    })


@dashboard_staff_required
def dashboard_email_settings(request):
    return render(request, 'dashboard/email_settings.html', {
        'active': 'email_settings',
        'smtp_host': settings.EMAIL_HOST,
        'smtp_user': settings.EMAIL_HOST_USER,
        'notify_email': settings.LEAD_RECIPIENT_EMAIL,
    })


@dashboard_staff_required
def dashboard_email_settings_test(request):
    if request.method == 'POST':
        try:
            send_store_email(
                'EduTrellis Store — test email',
                'This is a test email from your store\'s Email Settings page. If you received this, SMTP is configured correctly.',
                [get_notify_email()],
            )
            messages.success(request, f'Test email sent to {get_notify_email()}.')
        except Exception as exc:
            logger.exception("Test email failed: %s", exc)
            messages.error(request, f'Could not send test email: {exc}')
    return redirect('dashboard_email_settings')


@dashboard_staff_required
def dashboard_pwa_settings(request):
    settings_obj = PWASettings.get_solo()
    form = PWASettingsForm(request.POST or None, request.FILES or None, instance=settings_obj)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = PWASettingsForm(instance=settings_obj)
    return render(request, 'dashboard/pwa_settings.html', {
        'active': 'pwa_settings', 'form': form, 'settings_obj': settings_obj, 'saved': saved,
    })


@dashboard_staff_required
def dashboard_customize(request):
    settings_obj = SiteCustomization.get_solo()
    form = SiteCustomizationForm(request.POST or None, request.FILES or None, instance=settings_obj)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = SiteCustomizationForm(instance=settings_obj)
    return render(request, 'dashboard/customize.html', {
        'active': 'customize', 'form': form, 'settings_obj': settings_obj, 'saved': saved,
    })


@dashboard_staff_required
def dashboard_fee_settings(request):
    settings_obj = FeeSettings.get_solo()
    form = FeeSettingsForm(request.POST or None, instance=settings_obj)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = FeeSettingsForm(instance=settings_obj)
    return render(request, 'dashboard/fee_settings.html', {
        'active': 'fee_settings', 'form': form, 'settings_obj': settings_obj, 'saved': saved,
    })


@dashboard_staff_required
def dashboard_payments(request):
    q = request.GET.get('q', '').strip()
    status = request.GET.get('status', '').strip()
    payments = Payment.objects.select_related('order', 'order__user').order_by('-created_at')
    if q:
        payments = payments.filter(
            Q(order__id__icontains=q) | Q(order__user__username__icontains=q) |
            Q(order__user__email__icontains=q) | Q(razorpay_order_id__icontains=q) |
            Q(razorpay_payment_id__icontains=q)
        )
    if status:
        payments = payments.filter(status=status)
    return render(request, 'dashboard/payments.html', {
        'active': 'payments', 'payments': payments, 'q': q, 'status': status,
        'status_choices': Payment.STATUS_CHOICES,
    })


@dashboard_staff_required
def dashboard_backup(request):
    settings_obj = DropboxSettings.get_solo()
    backups = []
    list_error = None
    if settings_obj.is_configured:
        try:
            # Dropbox's API always returns client_modified as a naive UTC
            # datetime — Django's |date template filter only auto-converts
            # timezone-*aware* values to the local (IST) timezone, so left
            # naive this rendered as raw UTC clock time mislabeled as local.
            backups = [
                {'name': f.name, 'client_modified': timezone.localtime(f.client_modified.replace(tzinfo=dt_timezone.utc))}
                for f in dropbox_backup.list_backups(settings_obj)
            ]
        except dropbox_backup.BackupError as exc:
            list_error = str(exc)

    return render(request, 'dashboard/backup.html', {
        'active': 'backup', 'settings_obj': settings_obj, 'backups': backups,
        'list_error': list_error, 'dropbox_installed': dropbox_backup.dropbox is not None,
        'backup_folder': dropbox_backup.BACKUP_FOLDER,
    })


@dashboard_staff_required
def dashboard_backup_settings(request):
    settings_obj = DropboxSettings.get_solo()
    form = DropboxSettingsForm(request.POST or None, instance=settings_obj)
    saved = False
    if request.method == 'POST' and form.is_valid():
        form.save()
        saved = True
        form = DropboxSettingsForm(instance=settings_obj)
    return render(request, 'dashboard/backup_settings.html', {
        'active': 'backup', 'form': form, 'settings_obj': settings_obj, 'saved': saved,
        'dropbox_installed': dropbox_backup.dropbox is not None,
    })


@dashboard_staff_required
def dashboard_backup_run(request):
    if request.method == 'POST':
        settings_obj = DropboxSettings.get_solo()
        try:
            filename = dropbox_backup.create_backup(settings_obj)
            messages.success(request, f'Backup saved to Dropbox as "{filename}".')
        except dropbox_backup.BackupError as exc:
            messages.error(request, str(exc))
    return redirect('dashboard_backup')


@dashboard_staff_required
def dashboard_backup_restore(request):
    if request.method == 'POST':
        settings_obj = DropboxSettings.get_solo()
        filename = request.POST.get('filename', '').strip()
        if not filename:
            messages.error(request, 'Choose a backup to restore first.')
        else:
            try:
                dropbox_backup.restore_backup(settings_obj, filename)
                # Backups may have been created before newer application
                # migrations existed. Upgrade the restored schema before
                # redirecting to any page that imports the current models.
                call_command('migrate', interactive=False, verbosity=0)
                # restore_backup just swapped out db.sqlite3 from under this
                # very request, taking the django_session table — and this
                # request's own session row — with it. Without recreating it,
                # SessionMiddleware's save() at the end of the request can't
                # find the row to UPDATE and raises SessionInterrupted.
                # must_create=True forces an INSERT into the freshly-restored
                # table instead, so the admin doesn't get logged out or hit
                # an error page by restoring a backup.
                request.session.save(must_create=True)
                messages.success(request, f'Database restored from "{filename}" and upgraded to the current schema.')
            except dropbox_backup.BackupError as exc:
                messages.error(request, str(exc))
            except Exception:
                logger.exception('Restored database could not be migrated to the current schema')
                messages.error(
                    request,
                    'The database was restored, but its schema could not be upgraded. Run migrations before using the site.',
                )
    return redirect('dashboard_backup')


@dashboard_staff_required
def dashboard_backup_delete_all(request):
    if request.method == 'POST':
        confirmation = request.POST.get('confirmation', '').strip()
        if confirmation != 'DELETE ALL':
            messages.error(request, 'Type DELETE ALL exactly to confirm deleting every backup.')
            return redirect('dashboard_backup')
        settings_obj = DropboxSettings.get_solo()
        try:
            deleted = dropbox_backup.delete_all_backups(settings_obj)
            messages.success(request, f'Deleted {deleted} Dropbox backup item{"s" if deleted != 1 else ""}.')
        except dropbox_backup.BackupError as exc:
            messages.error(request, str(exc))
    return redirect('dashboard_backup')


def dashboard_logout(request):
    logout(request)
    return redirect('ai_page')


# ── AI Chat (/AI/) ──────────────────────────────────────────────────────────

AI_CHAT_RATE_LIMIT = 30           # messages
AI_CHAT_RATE_WINDOW = 10 * 60     # per 10 minutes, per IP
AI_CHAT_MAX_MESSAGE_CHARS = 16000
AI_CHAT_MAX_HISTORY = 20          # last 10 user+assistant turns — outer cap on how many rows are even fetched
# A per-message-count cap alone doesn't bound size: an attached document can
# replay up to 15,000 chars on every later turn, so a handful of document
# turns can approach the model's real context window even within 20
# messages. This is a second, size-based trim applied on top of the count
# cap (see the clean_history loop below) — roughly 4 chars/token, so this
# budget leaves headroom under typical 32k+ context windows once the system
# prompt, late reminders, and reply tokens are also accounted for.
AI_CHAT_HISTORY_CHAR_BUDGET = 48000
AI_CONVERSATION_TITLE_CHARS = 60
AI_CURRENT_CONVERSATION_SESSION_KEY = 'ai_current_conversation_id'
AI_GUEST_MESSAGE_LIMIT = 6        # free messages before a guest must log in/sign up
AI_FREE_MESSAGE_LIMIT = 20        # free messages for a logged-in, non-staff, unsubscribed account before Vidhyora AI requires the paid plan
AI_FREE_MODEL_KEYS = frozenset({'quick', 'code'})
# ~1.5MB of raw image data as a base64 data: URI (~2M chars) — well under
# Django's default 2.5MB DATA_UPLOAD_MAX_MEMORY_SIZE for the whole request
# body, so an oversized image gets our own clean error instead of Django's
# generic one. The client also resizes/compresses before ever uploading.
AI_IMAGE_MAX_DATA_URI_CHARS = 2_000_000
AI_DOCUMENT_MODES = {'coding', 'details'}
AI_DOCUMENT_CODE_MAX_OUTPUT_TOKENS = 6000
AI_IMAGE_GEN_HOURLY_LIMIT = 10       # FLUX generate/edit calls
AI_IMAGE_GEN_HOURLY_WINDOW = 60 * 60      # per 1 hour, per user (or per IP for guests)
AI_IMAGE_GEN_LOCKOUT_SECONDS = 6 * 60 * 60  # going over the hourly limit locks image generation out entirely for this long — harsher than the general chat limiter, which just makes you wait out the same window, since FLUX calls are the most expensive thing this endpoint does


def _ai_document_instruction(mode, filename, truncated=False):
    """Return the server-controlled instruction for an attachment action."""
    if mode == 'coding':
        truncation_rule = (
            " The supplied source was truncated, so clearly say that a complete safe rewrite is not possible "
            "from this partial input; do not pretend omitted sections are unchanged."
            if truncated else
            " Return the COMPLETE updated file, including unchanged sections, in one fenced code block; "
            "never return only a patch, diff, excerpt, or isolated snippet."
        )
        return (
            f"The user selected Start coding for the attached file {filename!r}. Apply exactly the change "
            f"requested in their current message.{truncation_rule} Keep explanation brief and put the full "
            "updated file first. For a binary office file, return the complete revised textual content that "
            "can be represented in chat and do not claim to have generated a downloadable binary file."
        )
    if mode == 'details':
        return (
            f"The user selected Show details for the attached file {filename!r}. Analyse and explain only: "
            "summarise its content and structure and point out relevant findings. Do not rewrite the file, "
            "do not output an updated version, and do not switch into coding unless the user asks in a later turn."
        )
    return None


AI_GENERATED_FILE_EXTENSIONS = frozenset({
    'txt', 'md', 'docx', 'html', 'htm', 'css', 'js', 'mjs', 'cjs', 'ts', 'tsx',
    'jsx', 'py', 'json', 'csv', 'xml', 'yaml', 'yml', 'sql', 'sh', 'ps1',
    'java', 'c', 'cpp', 'h', 'hpp', 'cs', 'php', 'rb', 'go', 'rs', 'swift',
    'kt', 'kts', 'vue', 'svelte', 'toml', 'ini',
})
AI_GENERATED_FILE_TYPE_EXTENSIONS = {
    'microsoft word': 'docx', 'word': 'docx', 'docx': 'docx',
    'markdown': 'md', 'html': 'html', 'css': 'css', 'javascript': 'js',
    'typescript': 'ts', 'python': 'py', 'json': 'json', 'csv': 'csv',
    'xml': 'xml', 'yaml': 'yaml', 'sql': 'sql', 'powershell': 'ps1',
    'java': 'java', 'php': 'php', 'ruby': 'rb', 'golang': 'go', 'rust': 'rs',
    'swift': 'swift', 'kotlin': 'kt', 'vue': 'vue', 'svelte': 'svelte',
}
_AI_FILE_ACTION_RE = re.compile(
    r"\b(?:create|generate|make|write|produce|prepare|build|export|save)\b",
    re.IGNORECASE,
)
_AI_FILE_OBJECT_RE = re.compile(
    r"\b(?:file|document|downloadable|download\s+link)\b",
    re.IGNORECASE,
)
_AI_EXPLICIT_FILENAME_RE = re.compile(
    r"\b([A-Za-z0-9][A-Za-z0-9_.-]{0,100}\.([A-Za-z0-9]{1,8}))\b",
    re.IGNORECASE,
)


def _ai_generated_file_spec(message):
    """Return a safe filename when the user explicitly asks for a file.

    This intent is handled server-side so every model-picker choice behaves
    consistently and none can pretend a download exists when it does not.
    """
    text = (message or '').strip()
    explicit = None
    for match in _AI_EXPLICIT_FILENAME_RE.finditer(text):
        if match.group(2).lower() in AI_GENERATED_FILE_EXTENSIONS:
            explicit = match.group(1)
            break
    if not _AI_FILE_ACTION_RE.search(text) or not (explicit or _AI_FILE_OBJECT_RE.search(text)):
        return None

    if explicit:
        candidate = explicit
    else:
        lower_text = text.lower()
        extension = next(
            (ext for label, ext in AI_GENERATED_FILE_TYPE_EXTENSIONS.items() if re.search(rf"\b{re.escape(label)}\b", lower_text)),
            'txt',
        )
        candidate = f'generated.{extension}'

    # Basename-only, conservative characters, and an allow-listed extension:
    # the result can be downloaded but never becomes a filesystem path.
    candidate = Path(candidate.replace('\\', '/')).name
    candidate = re.sub(r'[^A-Za-z0-9_.-]+', '-', candidate).strip('.-')[:120]
    stem, dot, extension = candidate.rpartition('.')
    if not dot or extension.lower() not in AI_GENERATED_FILE_EXTENSIONS:
        candidate = 'generated.txt'
    return {'file_name': candidate}


def _ai_generated_file_instruction(filename):
    if filename.lower().endswith('.docx'):
        return (
            f"The user explicitly requested a real downloadable Microsoft Word document named {filename!r}. "
            "Write the complete final document content now in exactly one fenced Markdown block. Use clear "
            "headings and lists where helpful, with no placeholders or explanatory text outside the block. "
            "Do not output XML, base64, a fake URL, or claim that you attached a file. The application will "
            "convert this content into a genuine DOCX file and attach the real download link."
        )
    return (
        f"The user explicitly requested a downloadable file named {filename!r}. Create the complete final "
        "contents now. Return exactly one fenced code block containing the entire file and no placeholders, "
        "patch, diff, download URL, attachment claim, or explanatory text outside that block. Preserve any "
        "exact content the user supplied. The application, not you, will create and attach the real download link."
    )


def _extract_ai_generated_file_content(reply):
    """Extract the complete fenced payload requested by the server prompt."""
    text = (reply or '').strip()
    fenced = re.search(r"```[^\r\n]*\r?\n([\s\S]*?)```", text)
    if fenced:
        return fenced.group(1).rstrip('\r\n')
    if text.startswith('```') and text.endswith('```'):
        text = text[3:-3]
        text = re.sub(r'^[A-Za-z0-9_+.-]+\r?\n', '', text, count=1)
    return text.strip()


def _ai_word_document_bytes(content):
    """Convert model-produced document text into a genuine DOCX package."""
    document = WordDocument()
    for raw_line in (content or '').splitlines():
        line = raw_line.strip()
        heading = re.match(r'^(#{1,6})\s+(.+)$', line)
        bullet = re.match(r'^[-*]\s+(.+)$', line)
        numbered = re.match(r'^\d+[.)]\s+(.+)$', line)
        if heading:
            document.add_heading(heading.group(2), level=min(len(heading.group(1)), 9))
        elif bullet:
            document.add_paragraph(bullet.group(1), style='List Bullet')
        elif numbered:
            document.add_paragraph(numbered.group(1), style='List Number')
        else:
            document.add_paragraph(raw_line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _ai_user_context(user):
    """A compact snapshot of the logged-in user's own remembered name/
    location — nothing about the store (no cart, wallet, or order data;
    the AI chat no longer has any access to that) — so the assistant can
    address a returning user naturally instead of asking who they are
    every conversation. Always scoped to `user` (the authenticated
    request.user)."""
    lines = [f"Logged in as: {user.first_name or user.username}."]

    profile = getattr(user, 'store_profile', None)
    if profile:
        # Name/location the user gave the AI chat during onboarding (see
        # ai_chat_send's onboarding block) — surfaced here so a returning
        # user is actually remembered/addressed naturally instead of the
        # assistant asking who they are again every conversation.
        if profile.ai_display_name:
            lines.append(f"This user previously told the assistant their name is {profile.ai_display_name}.")
        if profile.ai_location:
            lines.append(f"This user previously told the assistant they're located in {profile.ai_location}.")

    return '\n'.join(lines)


def _ai_owner_filter(request):
    """Logged-in users own conversations by user FK; guests own them by the
    session_key their browser already carries (same session used for the
    anonymous cart) — that's what lets a guest's chat survive a page reload
    and then get handed to their account the moment they log in. Deliberately
    does NOT create a session for a guest that doesn't have one yet — a plain
    page view or list/read call has no conversation to attach to anyway, and
    forcing a new database-backed session row on every cookie-less visit is
    an easy way to grow django_session unbounded. Only ai_chat_send (which
    actually needs a stable key to save a message under) creates one."""
    if request.user.is_authenticated:
        return Q(user=request.user)
    session_key = request.session.session_key
    if not session_key:
        return Q(pk__in=[])
    return Q(user__isnull=True, session_key=session_key)


def ai_generated_file_download(request, token):
    """Download one AI-created file, only for the account/session that made it."""
    if request.method != 'GET':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    files = AIGeneratedFile.objects.filter(token=token)
    if request.user.is_authenticated:
        files = files.filter(user=request.user)
    else:
        session_key = request.session.session_key
        if not session_key:
            files = files.none()
        else:
            files = files.filter(user__isnull=True, session_key=session_key)
    generated_file = get_object_or_404(files)

    if generated_file.file_name.lower().endswith('.docx'):
        payload = _ai_word_document_bytes(generated_file.content)
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        payload = generated_file.content.encode('utf-8')
        content_type = f"{mimetypes.guess_type(generated_file.file_name)[0] or 'text/plain'}; charset=utf-8"
    response = HttpResponse(payload, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{generated_file.file_name}"'
    response['X-Content-Type-Options'] = 'nosniff'
    response['Cache-Control'] = 'private, no-store'
    return response


def _ai_notes_snapshot(request):
    """Single database snapshot used by page boot, API responses and chat."""
    notes = list(
        AINote.objects.filter(_ai_owner_filter(request))
        .values('id', 'heading', 'content', 'created_at')
        .order_by('-created_at', '-pk')[:200]
    )
    for note in notes:
        note['created_at'] = timezone.localtime(note['created_at']).isoformat()
    return notes


def ai_manifest(request):
    pwa = PWASettings.get_solo()
    version = int(pwa.updated_at.timestamp() * 1_000_000)
    if pwa.icon:
        icon_192 = request.build_absolute_uri(
            f"{reverse('ai_pwa_icon', args=[192])}?v={version}"
        )
        icon_512 = request.build_absolute_uri(
            f"{reverse('ai_pwa_icon', args=[512])}?v={version}"
        )
    else:
        icon_192 = request.build_absolute_uri(static_url('ai-icon-192.png'))
        icon_512 = request.build_absolute_uri(static_url('ai-icon-512.png'))
    manifest = {
        'id': '/',
        'name': pwa.app_name or 'Vidhyora AI',
        'short_name': pwa.short_name or 'Vidhyora AI',
        'description': pwa.description or 'Chat with Vidhyora AI.',
        'start_url': '/',
        'scope': '/',
        'display': 'standalone',
        'background_color': pwa.background_color or '#ffffff',
        'theme_color': pwa.theme_color or '#059669',
        'icons': [
            {'src': icon_192, 'sizes': '192x192', 'type': 'image/png', 'purpose': 'any'},
            {'src': icon_512, 'sizes': '512x512', 'type': 'image/png', 'purpose': 'any'},
            {'src': icon_512, 'sizes': '512x512', 'type': 'image/png', 'purpose': 'maskable'},
        ],
    }
    response = JsonResponse(manifest, content_type='application/manifest+json')
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return response


def ai_pwa_icon(request, size):
    """Serve the uploaded admin icon at real manifest-required dimensions."""
    if size not in (192, 512):
        return JsonResponse({'status': 'error', 'detail': 'Invalid icon size.'}, status=404)
    pwa = PWASettings.get_solo()
    if not pwa.icon:
        return JsonResponse({'status': 'error', 'detail': 'No PWA icon configured.'}, status=404)
    try:
        pwa.icon.open('rb')
        with Image.open(pwa.icon) as source:
            source = ImageOps.exif_transpose(source).convert('RGBA')
            rendered = ImageOps.fit(
                source, (size, size), method=Image.Resampling.LANCZOS,
            )
            output = io.BytesIO()
            rendered.save(output, format='PNG', optimize=True)
        pwa.icon.close()
    except Exception:
        logger.exception('Could not render configured PWA icon')
        return JsonResponse({'status': 'error', 'detail': 'Could not load the PWA icon.'}, status=404)
    response = HttpResponse(output.getvalue(), content_type='image/png')
    response['Cache-Control'] = 'public, max-age=3600'
    return response


def ai_page(request):
    pwa = PWASettings.get_solo()
    pwa_version = int(pwa.updated_at.timestamp() * 1_000_000)
    conversations = list(
        AIConversation.objects.filter(_ai_owner_filter(request))
        .values('id', 'title', 'updated_at').order_by('-updated_at')[:100]
    )
    for c in conversations:
        c['updated_at'] = timezone.localtime(c['updated_at']).isoformat()

    notes = _ai_notes_snapshot(request)

    # Browser storage can be unavailable or cleared. Remember the last chat
    # the server actually opened/sent to as a safe refresh/login fallback.
    # Always validate it against this request's owner before exposing it.
    conversation_ids = {c['id'] for c in conversations}
    try:
        resume_conversation_id = int(request.session.get(AI_CURRENT_CONVERSATION_SESSION_KEY) or 0)
    except (TypeError, ValueError):
        resume_conversation_id = 0
    if resume_conversation_id not in conversation_ids:
        resume_conversation_id = conversations[0]['id'] if conversations else None
    if resume_conversation_id:
        request.session[AI_CURRENT_CONVERSATION_SESSION_KEY] = resume_conversation_id
    model_labels = {key: cfg['label'] for key, cfg in ai_chat.MODELS.items()}

    ai_is_staff = bool(request.user.is_authenticated and request.user.is_staff)
    ai_subscribed = False
    ai_free_used = 0
    if request.user.is_authenticated and not ai_is_staff:
        profile, _ = StoreProfile.objects.get_or_create(user=request.user)
        ai_subscribed = profile.is_ai_subscribed
        ai_free_used = profile.ai_free_messages_used
    ai_full_model_access = bool(ai_is_staff or ai_subscribed)
    ai_default_model = ai_chat.DEFAULT_MODEL_KEY if ai_full_model_access else 'quick'
    models = [
        {
            'key': key,
            'label': cfg['label'],
            'description': cfg['description'],
            'locked': not ai_full_model_access and key not in AI_FREE_MODEL_KEYS,
        }
        for key, cfg in ai_chat.MODELS.items() if key != 'vision'
    ]

    return render(request, 'ai.html', {
        'ai_authenticated': request.user.is_authenticated,
        'ai_user': _user_payload(request.user) if request.user.is_authenticated else None,
        'ai_conversations': conversations,
        'ai_notes': notes,
        'ai_resume_conversation_id': resume_conversation_id,
        'ai_guest_limit': AI_GUEST_MESSAGE_LIMIT,
        'ai_guest_used': 0 if request.user.is_authenticated else min(
            AI_GUEST_MESSAGE_LIMIT,
            max(request.session.get('ai_guest_msg_count', 0), _ip_free_messages_used(_client_ip(request))),
        ),
        'ai_is_staff': ai_is_staff,
        'ai_subscribed': ai_subscribed,
        'ai_free_limit': AI_FREE_MESSAGE_LIMIT,
        'ai_free_used': ai_free_used,
        'ai_purchase_url': _ai_purchase_url(),
        'ai_models': models,
        'ai_default_model': ai_default_model,
        'ai_default_model_label': model_labels[ai_default_model],
        'ai_model_labels': model_labels,
        'ai_github_oauth_available': bool(settings.GITHUB_OAUTH_CLIENT_ID),
        'show_location_prompt': _location_prompt_needed(request.user),
        'show_profile_wizard': _profile_wizard_needed(request.user),
        'ai_pwa': pwa,
        'ai_pwa_ready': pwa.ready,
        'ai_pwa_version': pwa_version,
    })


def _client_ip(request):
    forwarded = request.META.get('HTTP_X_FORWARDED_FOR')
    return forwarded.split(',')[0].strip() if forwarded else request.META.get('REMOTE_ADDR', 'unknown')


def _format_wait_time(seconds):
    """'42 seconds' under a minute, '3 minutes' under an hour, otherwise
    '6 hours' — each rounded up to the next whole unit, so the number shown
    is never an underestimate of how long is actually left."""
    seconds = max(1, int(seconds))
    if seconds < 60:
        return f"{seconds} second{'s' if seconds != 1 else ''}"
    if seconds < 3600:
        minutes = (seconds + 59) // 60
        return f"{minutes} minute{'s' if minutes != 1 else ''}"
    hours = (seconds + 3599) // 3600
    return f"{hours} hour{'s' if hours != 1 else ''}"


def _ai_has_full_model_access(user):
    if not user.is_authenticated:
        return False
    if user.is_staff:
        return True
    profile, _ = StoreProfile.objects.get_or_create(user=user)
    return profile.is_ai_subscribed


def _ai_purchase_url():
    # Premium access is managed by staff from the AI Management dashboard.
    return 'mailto:support@edutrellis.in?subject=Vidhyora%20AI%20Premium'


def _ip_free_messages_used(ip):
    """Total free (guest + signed-in-but-unsubscribed) Vidhyora AI
    messages ever sent from this IP, across every guest session and every
    account that's ever chatted from it. Session/account counters alone
    (request.session['ai_guest_msg_count'], StoreProfile.ai_free_messages_
    used) reset the moment someone opens an incognito window or signs up a
    throwaway second account — this is what actually survives that, since
    it's derived from the saved messages themselves, not a counter. Deliber-
    ately counts every non-staff message tied to this IP regardless of
    whether the sender was subscribed at the time — an over-count is the
    safe direction for an anti-abuse cap."""
    if not ip or ip == 'unknown':
        return 0
    return AIMessage.objects.filter(
        role=AIMessage.ROLE_USER, conversation__ip_address=ip,
    ).exclude(conversation__user__is_staff=True).count()


def _ai_profile_gate(user, ip=None):
    """Returns None if `user` (already known to be authenticated) has
    unrestricted Vidhyora AI access — staff, or an active paid
    subscription — else the 403 JSON payload to send back once they've used
    their free allotment. Staff never even touch the StoreProfile row here,
    which is what gives every staff account (including admin@gmail.com)
    unlimited messages and every model with no separate per-account
    allowlist to maintain."""
    if user.is_staff:
        return None
    profile, _ = StoreProfile.objects.get_or_create(user=user)
    if profile.is_ai_subscribed:
        return None
    ip_capped = bool(ip) and _ip_free_messages_used(ip) >= (AI_GUEST_MESSAGE_LIMIT + AI_FREE_MESSAGE_LIMIT)
    if profile.ai_free_messages_used >= AI_FREE_MESSAGE_LIMIT or ip_capped:
        return {
            'status': 'subscription_required',
            'detail': (
                f"You've used all {AI_FREE_MESSAGE_LIMIT} free Vidhyora AI messages. "
                "Subscribe for Rs 99/month for unlimited messages on every model."
            ),
            'purchase_url': _ai_purchase_url(),
        }
    return None


def _ai_note_heading(text):
    """First line of the noted text, trimmed to a short heading — same idea
    as AIConversation's title-from-first-message, just capped shorter since
    this is meant to read like a Google Keep card title."""
    first_line = text.strip().splitlines()[0].strip() if text.strip() else ''
    return (first_line[:60] + '…') if len(first_line) > 60 else first_line


AI_SELECTED_NOTES_SESSION_KEY = 'ai_selected_note_ids'
AI_PENDING_NOTE_EDITS_SESSION_KEY = 'ai_pending_note_edits'


def _ai_safe_note_text(text):
    """Fix only a small set of unambiguous spelling mistakes."""
    corrections = {
        'teh': 'the', 'tomorow': 'tomorrow', 'tommorow': 'tomorrow',
        'meting': 'meeting', 'remeber': 'remember',
    }
    return re.sub(
        r"\b(?:teh|tomorow|tommorow|meting|remeber)\b",
        lambda match: corrections[match.group(0).lower()],
        (text or '').strip(), flags=re.IGNORECASE,
    )


def _ai_set_selected_note(request, conversation, note_id):
    selected = dict(request.session.get(AI_SELECTED_NOTES_SESSION_KEY, {}))
    selected[str(conversation.pk)] = int(note_id)
    request.session[AI_SELECTED_NOTES_SESSION_KEY] = selected


def _ai_get_selected_note(request, conversation):
    selected = request.session.get(AI_SELECTED_NOTES_SESSION_KEY, {})
    try:
        note_id = int(selected.get(str(conversation.pk)))
    except (TypeError, ValueError, AttributeError):
        return None
    return AINote.objects.filter(_ai_owner_filter(request), pk=note_id).first()


def _ai_set_pending_note_edit(request, conversation, payload=None):
    pending = dict(request.session.get(AI_PENDING_NOTE_EDITS_SESSION_KEY, {}))
    key = str(conversation.pk)
    if payload:
        pending[key] = payload
    else:
        pending.pop(key, None)
    request.session[AI_PENDING_NOTE_EDITS_SESSION_KEY] = pending


def _ai_get_pending_note_edit(request, conversation):
    pending = request.session.get(AI_PENDING_NOTE_EDITS_SESSION_KEY, {})
    return pending.get(str(conversation.pk)) if isinstance(pending, dict) else None


def _ai_note_action_reply(request, conversation, confirmation, extra_headers=None):
    """Shared tail for every My Notes action (save/show/delete/edit) —
    saves `confirmation` as the assistant's turn and returns it on the same
    StreamingHttpResponse contract ai_chat_send's real model-call path uses,
    so the frontend's existing send()/pump() handling needs no
    special-casing beyond reading the extra X-Notes-Changed header (see
    refreshNotes() in ai.html). No AI model call, no free-message quota
    spent, for any of these."""
    AIMessage.objects.create(
        conversation=conversation, role=AIMessage.ROLE_ASSISTANT,
        content=confirmation, model_key='note',
    )
    conversation.updated_at = timezone.now()
    conversation.save(update_fields=['updated_at'])

    def event_stream():
        yield confirmation

    response = StreamingHttpResponse(event_stream(), content_type='text/plain; charset=utf-8')
    response['Cache-Control'] = 'no-cache'
    response['X-Accel-Buffering'] = 'no'
    response['X-Conversation-Id'] = str(conversation.id)
    response['X-Model-Key'] = 'note'
    response['X-Request-Category'] = 'note'
    response['X-Sumudrika'] = ''
    response['X-Jagu'] = ''
    response['X-Persona-End'] = ''
    for key, value in (extra_headers or {}).items():
        response[key] = value
    return response


def _ai_save_note_response(request, conversation, message):
    """'Take this note' / 'note it down' / 'save details...' (see
    request_router.is_note_intent) — saves a new AINote.

    Notes whatever extra text came with the trigger phrase itself ('note
    down: buy milk' -> 'buy milk') when there is any; otherwise falls back
    to the assistant's last reply in this conversation, since that's what a
    bare 'note it down' naturally refers to — and to the user's own message
    as a last resort, for a brand new conversation with nothing to look
    back at yet.
    """
    remainder = request_router.strip_note_intent(message).strip(' :-—')
    if not remainder:
        return _ai_note_action_reply(request, conversation, 'What would you like the note to say?')
    note_content = _ai_safe_note_text(remainder)
    heading = _ai_note_heading(note_content)

    note = AINote.objects.create(
        user=request.user if request.user.is_authenticated else None,
        session_key='' if request.user.is_authenticated else (request.session.session_key or ''),
        conversation=conversation, heading=heading, content=note_content,
    )

    note = AINote.objects.filter(_ai_owner_filter(request), pk=note.pk).first()
    if not note:
        return _ai_note_action_reply(request, conversation, 'I couldn\'t save that note. Please try again.')
    _ai_notes_snapshot(request)
    _ai_set_selected_note(request, conversation, note.pk)

    confirmation = "Saved to your notes. The saved note is:\n\n" + _ai_note_final_text(note)
    return _ai_note_action_reply(request, conversation, confirmation, {
        'X-Notes-Changed': '1', 'X-Note-Id': str(note.id),
    })


def _ai_show_notes_response(request, conversation):
    """'Show my notes' / 'what notes do I have' (see
    request_router.is_show_notes_intent) — lists the user's saved notes
    right in the chat, newest first."""
    notes = _ai_notes_snapshot(request)
    if not notes:
        confirmation = "You don’t have any saved notes."
    else:
        lines = [
            (f"{index}. **{note['content']}**" if (note['heading'] or '').strip() == note['content'].strip()
             else f"{index}. **{note['heading'] or 'Untitled note'}**\n{note['content']}")
            for index, note in enumerate(notes, 1)
        ]
        confirmation = "Here are your saved notes:\n\n" + '\n\n'.join(lines)
    return _ai_note_action_reply(request, conversation, confirmation)


def _ai_matching_notes(request, target):
    database_id = re.fullmatch(r"id\s*#?\s*(\d+)", (target or '').strip(), re.IGNORECASE)
    if database_id:
        note = AINote.objects.filter(_ai_owner_filter(request), pk=int(database_id.group(1))).first()
        return [note] if note else []
    ordinal = re.fullmatch(r"(?:number\s*|#\s*)?(\d+)(?:st|nd|rd|th)?", (target or '').strip(), re.IGNORECASE)
    if ordinal:
        position = int(ordinal.group(1))
        if position < 1:
            return []
        note = AINote.objects.filter(_ai_owner_filter(request)).order_by('-created_at', '-pk')[position - 1:position].first()
        return [note] if note else []
    return list(AINote.objects.filter(_ai_owner_filter(request)).filter(
        Q(heading__icontains=target) | Q(content__icontains=target),
    )[:6])


def _ai_read_note_response(request, conversation, target):
    """Open one owned note by its displayed number or a unique title phrase."""
    if not target:
        return _ai_note_action_reply(request, conversation, 'Tell me which note to open, e.g. "open note 1".')
    matches = _ai_matching_notes(request, target)
    if not matches:
        return _ai_note_action_reply(request, conversation, f'I couldn\'t find a note matching "{target}".')
    if len(matches) > 1:
        listing = '\n'.join(f"- {n.heading or 'Untitled note'}" for n in matches)
        return _ai_note_action_reply(request, conversation, f"I found more than one matching note:\n\n{listing}\n\nUse its number from \"show my notes\" or a more specific title.")
    note = matches[0]
    _ai_set_selected_note(request, conversation, note.pk)
    created = timezone.localtime(note.created_at).strftime('%b %d, %Y at %I:%M %p')
    return _ai_note_action_reply(
        request, conversation,
        f"**{note.heading or 'Untitled note'}**\n\n{note.content}\n\n_Saved {created}_",
        {'X-Note-Id': str(note.id)},
    )


def _ai_delete_note_response(request, conversation, target):
    """'Delete note about X' / 'delete all my notes' (see
    request_router.match_delete_note)."""
    if target == request_router.DELETE_ALL_NOTES:
        owner_notes = AINote.objects.filter(_ai_owner_filter(request))
        count = owner_notes.count()
        owner_notes.delete()
        if _ai_notes_snapshot(request):
            return _ai_note_action_reply(request, conversation, 'I couldn\'t delete all notes. Please try again.')
        confirmation = f"Deleted all {count} of your saved notes." if count else "You don’t have any saved notes."
        return _ai_note_action_reply(request, conversation, confirmation, {'X-Notes-Changed': '1'})

    if not target:
        confirmation = 'Tell me which note to delete, e.g. "delete note about milk".'
        return _ai_note_action_reply(request, conversation, confirmation)

    matches = _ai_matching_notes(request, target)
    if not matches:
        confirmation = f'I couldn\'t find a note matching "{target}".'
        return _ai_note_action_reply(request, conversation, confirmation)
    if len(matches) > 1:
        listing = '\n'.join(f"- {n.heading or 'Untitled note'}" for n in matches)
        confirmation = f"I found more than one note matching that:\n\n{listing}\n\nBe more specific about which one to delete."
        return _ai_note_action_reply(request, conversation, confirmation)

    note = matches[0]
    heading = note.heading or 'Untitled note'
    note_id = note.pk
    note.delete()
    latest_notes = _ai_notes_snapshot(request)
    if any(item['id'] == note_id for item in latest_notes):
        return _ai_note_action_reply(request, conversation, 'I couldn\'t delete that note. Please try again.')
    confirmation = f"Deleted your note — **{heading}**."
    return _ai_note_action_reply(request, conversation, confirmation, {'X-Notes-Changed': '1'})


_AI_NOTE_TIME_RE = re.compile(
    r"\b(?:[01]?\d|2[0-3])(?::[0-5]\d)?\s*(?:am|pm)\b|"
    r"\b(?:[01]?\d|2[0-3]):[0-5]\d\b",
    re.IGNORECASE,
)


def _ai_note_final_text(note):
    if (note.heading or '').strip() == note.content.strip():
        return f"**{note.content}**"
    return f"**{note.heading or 'Untitled note'}**\n\n{note.content}"


def _ai_commit_note_update(request, conversation, note, *, content=None, heading=None):
    update_fields = []
    if content is not None:
        note.content = _ai_safe_note_text(content)
        note.heading = _ai_note_heading(note.content)
        update_fields.extend(['content', 'heading'])
    if heading is not None:
        note.heading = _ai_safe_note_text(heading)[:120]
        if 'heading' not in update_fields:
            update_fields.append('heading')
    note.save(update_fields=update_fields)

    saved = AINote.objects.filter(_ai_owner_filter(request), pk=note.pk).first()
    expected_content = note.content if content is not None else saved.content if saved else None
    expected_heading = note.heading if heading is not None else saved.heading if saved else None
    if not saved or saved.content != expected_content or saved.heading != expected_heading:
        return _ai_note_action_reply(request, conversation, 'I couldn\'t update that note. Please try again.')

    _ai_notes_snapshot(request)
    _ai_set_selected_note(request, conversation, saved.pk)
    _ai_set_pending_note_edit(request, conversation)
    confirmation = "Updated your note. The saved note is now:\n\n" + _ai_note_final_text(saved)
    return _ai_note_action_reply(request, conversation, confirmation, {
        'X-Notes-Changed': '1', 'X-Note-Id': str(saved.pk),
    })


def _ai_replace_note_time(request, conversation, note, new_time, old_time=None):
    if old_time:
        match = re.search(re.escape(old_time), note.content, re.IGNORECASE)
        if not match:
            return _ai_note_action_reply(request, conversation, f'I couldn\'t find "{old_time}" in that note.')
        updated = note.content[:match.start()] + new_time + note.content[match.end():]
    else:
        times = list(_AI_NOTE_TIME_RE.finditer(note.content))
        if not times:
            return _ai_note_action_reply(request, conversation, 'That note does not contain a time to update.')
        if len(times) > 1:
            choices = ', '.join(match.group(0) for match in times)
            return _ai_note_action_reply(request, conversation, f"That note contains multiple times ({choices}). Which one should I change?")
        match = times[0]
        updated = note.content[:match.start()] + new_time + note.content[match.end():]
    return _ai_commit_note_update(request, conversation, note, content=updated)


def _ai_apply_note_instruction(request, conversation, note, instruction, original_message):
    instruction = (instruction or '').strip(' :-—')
    original_message = original_message or instruction
    _ai_set_selected_note(request, conversation, note.pk)

    if not instruction:
        return _ai_note_action_reply(
            request, conversation,
            _ai_note_final_text(note) + "\n\nWhat would you like to change in this note?",
            {'X-Note-Id': str(note.pk)},
        )

    if re.search(r"\brename\b", original_message, re.IGNORECASE):
        return _ai_commit_note_update(request, conversation, note, heading=instruction)

    time_change = re.search(
        r"(?:(?:update|change|correct|set)\s+)?"
        r"(?:(?:only\s+)?(?:the\s+)?(?:meeting\s+)?time|(?:this|the)\s+note(?:'s)?\s+(?:meeting\s+)?time)"
        r"(?:\s+from\s+(?P<old>(?:[01]?\d|2[0-3])(?::[0-5]\d)?\s*(?:am|pm)))?"
        r"\s+to\s+(?P<new>(?:[01]?\d|2[0-3])(?::[0-5]\d)?\s*(?:am|pm))\b",
        original_message, re.IGNORECASE,
    )
    if time_change:
        return _ai_replace_note_time(
            request, conversation, note,
            time_change.group('new').strip(),
            (time_change.group('old') or '').strip() or None,
        )

    exact_change = re.search(
        r"(?:change|replace|correct)\s+(.+?)\s+(?:to|with)\s+(.+?)\s*$",
        original_message, re.IGNORECASE,
    )
    if exact_change and exact_change.group(1).lower().strip() not in ('it', 'this', 'this note', 'the note'):
        old, new = exact_change.group(1).strip(), _ai_safe_note_text(exact_change.group(2))
        found = re.search(re.escape(old), note.content, re.IGNORECASE)
        if found:
            updated = note.content[:found.start()] + new + note.content[found.end():]
            return _ai_commit_note_update(request, conversation, note, content=updated)

    explicit_full = bool(re.search(r"\b(?:replace|rewrite|full note|entire note)\b", original_message, re.IGNORECASE))
    if explicit_full:
        return _ai_commit_note_update(request, conversation, note, content=instruction)

    possible_time = re.fullmatch(r"(?:[01]?\d|2[0-3])(?::[0-5]\d)?\s*(?:am|pm)", instruction, re.IGNORECASE)
    field = 'time' if possible_time and _AI_NOTE_TIME_RE.search(note.content) else 'part'
    _ai_set_pending_note_edit(request, conversation, {
        'note_id': note.pk, 'new_content': instruction, 'field': field,
    })
    if field == 'time':
        question = f"Should I update only the time to {instruction}, or replace the full note?"
    else:
        question = "Should I update only part of this note, or replace the full note?"
    return _ai_note_action_reply(request, conversation, question, {'X-Note-Id': str(note.pk)})


def _ai_edit_note_response(request, conversation, target, new_content, original_message):
    if not target:
        return _ai_note_action_reply(request, conversation, 'Which note would you like to edit? Use its sidebar number or database ID.')

    matches = _ai_matching_notes(request, target)
    if not matches:
        return _ai_note_action_reply(request, conversation, f'I couldn\'t find a note matching "{target}".')
    if len(matches) > 1:
        listing = '\n'.join(f"- {n.heading or 'Untitled note'}" for n in matches)
        return _ai_note_action_reply(request, conversation, f"I found more than one matching note:\n\n{listing}\n\nWhich note do you mean?")

    note = matches[0]
    return _ai_apply_note_instruction(request, conversation, note, new_content, original_message)


def _ai_contextual_note_edit_response(request, conversation, message):
    note = _ai_get_selected_note(request, conversation)
    if not note:
        return None
    ambiguous = re.search(r"(?:update|change|set)\s+(?:it|this note|the note)\s+to\s+(.+?)\s*$", message, re.IGNORECASE)
    instruction = ambiguous.group(1).strip() if ambiguous else message
    return _ai_apply_note_instruction(request, conversation, note, instruction, message)


def _ai_pending_note_edit_response(request, conversation, message):
    pending = _ai_get_pending_note_edit(request, conversation)
    if not pending:
        return None
    note = AINote.objects.filter(_ai_owner_filter(request), pk=pending.get('note_id')).first()
    if not note:
        _ai_set_pending_note_edit(request, conversation)
        return _ai_note_action_reply(request, conversation, 'That note no longer exists.')
    if re.fullmatch(r"\s*(?:(?:only|just)\s+(?:the\s+)?time|update\s+only\s+(?:the\s+)?time)\s*", message, re.IGNORECASE):
        if pending.get('field') != 'time':
            return _ai_note_action_reply(request, conversation, 'What should the new time be?')
        return _ai_replace_note_time(request, conversation, note, pending['new_content'])
    if re.search(r"\b(?:replace|rewrite)\s+(?:the\s+)?(?:full|entire|whole)?\s*note\b|\bfull replacement\b", message, re.IGNORECASE):
        return _ai_commit_note_update(request, conversation, note, content=pending['new_content'])
    return None


def _ai_image_gen_gate(user, ip):
    """Rate-gates FLUX image generation/editing: AI_IMAGE_GEN_HOURLY_LIMIT
    calls per rolling AI_IMAGE_GEN_HOURLY_WINDOW, keyed by user id when
    logged in, else IP (same key shape as the general chat limiter above).
    Counted at attempt time, before generation runs, same reasoning as the
    general limiter — otherwise a blocked user could keep retrying for
    free. Going over the hourly limit locks image generation out entirely
    for AI_IMAGE_GEN_LOCKOUT_SECONDS rather than just making them wait out
    the same window. Returns a user-facing detail string when blocked,
    else None."""
    key = f'user:{user.id}' if user.is_authenticated else f'ip:{ip}'
    lockout_key = f'ai_image_lockout:{key}'
    now = time.time()
    lockout_until = cache.get(lockout_key)
    if lockout_until and now < lockout_until:
        return f"Image generation limit reached. Please try again in {_format_wait_time(lockout_until - now)}."

    count_key = f'ai_image_gen_count:{key}'
    window = cache.get(count_key)
    if not window or now >= window['reset_at']:
        window = {'count': 0, 'reset_at': now + AI_IMAGE_GEN_HOURLY_WINDOW}
    window['count'] += 1
    if window['count'] > AI_IMAGE_GEN_HOURLY_LIMIT:
        cache.set(lockout_key, now + AI_IMAGE_GEN_LOCKOUT_SECONDS, AI_IMAGE_GEN_LOCKOUT_SECONDS)
        return f"Image generation limit reached. Please try again in {_format_wait_time(AI_IMAGE_GEN_LOCKOUT_SECONDS)}."
    cache.set(count_key, window, AI_IMAGE_GEN_HOURLY_WINDOW)
    return None


def _chatgpt_image_error_detail(error):
    """Hide internal image providers/workers behind the ChatGPT identity."""
    detail = str(error).strip()
    lower_detail = detail.lower()
    if 'content filter' in lower_detail or 'content_filtered' in lower_detail:
        return 'That image request was blocked by the safety filter. Try a different prompt or image.'
    if 'limit' in lower_detail or getattr(error, 'status_code', None) == 429:
        return 'The image-generation limit has been reached. Please wait and try again later.'
    if any(term in lower_detail for term in (
        'nvidia', 'flux', 'nemotron', 'black-forest', 'black forest', 'api_key',
    )):
        return 'ChatGPT 5.6 could not generate that image. Try a different prompt or image.'
    return detail or 'ChatGPT 5.6 could not generate that image. Please try again.'


_CHATGPT_HIDDEN_MODEL_PATTERNS = (
    re.compile(
        r'black[-\s]?forest[-\s]?labs\s*/\s*flux(?:[_\s.-]*2)?(?:[_\s.-]*klein)?(?:[_\s.-]*4b)?',
        re.IGNORECASE,
    ),
    re.compile(r'\bflux(?:\s*\.?\s*2)?(?:[\s._-]*klein)?(?:[\s._-]*4b)?(?:\s+model)?\b', re.IGNORECASE),
    re.compile(
        r'\b(?:nvidia[\s/-]+)?nemotron(?:[-\s._]*3)?(?:[-\s._]*(?:nano|super|ultra|omni))?'
        r'(?:[-\s._]*\d+[a-z]?)?(?:[-\s._]*a\d+b)?(?:[-\s._]*reasoning)?(?:\s+model)?\b',
        re.IGNORECASE,
    ),
    re.compile(r'\bVidhyora\s+(?:Ultra|Quick|Light|Code|Vision)(?:\s+model)?\b', re.IGNORECASE),
)


def _chatgpt_public_reply(reply):
    """Keep routed worker identities out of ChatGPT-visible response text."""
    cleaned = str(reply or '')
    for pattern in _CHATGPT_HIDDEN_MODEL_PATTERNS:
        cleaned = pattern.sub('ChatGPT 5.6', cleaned)
    return cleaned


def _ai_public_routed_model_key(response_model_key, routed_model_key):
    """Never expose ChatGPT's private worker selection to the browser."""
    if response_model_key == ai_chat.CHATGPT_56_MODEL_KEY:
        return ai_chat.CHATGPT_56_MODEL_KEY
    return routed_model_key


def _ai_chat_failure_reply(error, response_model_key, is_staff=False):
    """Turn upstream failures into safe, accurate, public-facing guidance."""
    label = ai_chat.MODELS.get(response_model_key, {}).get('label', 'The selected AI model')
    status_code = getattr(error, 'status_code', None)
    if ai_chat._is_model_unavailable_error(error):
        suffix = (
            ' Update the configured text-model API key, then restart the application.'
            if is_staff else
            ' Please try again after the administrator reconnects it.'
        )
        return f'{label} text access is currently disconnected.{suffix}'
    if status_code in (401, 403):
        suffix = (
            ' Update the configured API key, then restart the application.'
            if is_staff else
            ' Please try again after the administrator reconnects it.'
        )
        return f'{label} authentication is currently unavailable.{suffix}'
    if status_code == 429:
        return f'{label} is currently at its request limit. Please wait a moment and try again.'
    if ai_chat._is_transient_error(error):
        return f'{label} is temporarily unavailable. Please wait a moment and try again.'
    return (
        f'{label} is taking longer than expected to respond. Please try again in '
        'a moment — if this keeps happening, switching to a different model from '
        'the picker usually helps.'
    )


def _ai_flux_response(conversation, prompt, source_image, response_model_key=None):
    """Run a FLUX generation/editing turn and persist the real image URL."""
    display_model_key = response_model_key or ai_chat.FLUX_KLEIN_4B_MODEL_KEY
    try:
        generated = image_generation.generate_image(prompt, source_image or None)
        dated_folder = timezone.now().strftime('%Y/%m/%d')
        filename = f"ai_generated/{dated_folder}/{secrets.token_hex(16)}.{generated.extension}"
        stored_name = default_storage.save(filename, ContentFile(generated.content))
        generated_url = default_storage.url(stored_name)
    except image_generation.ImageGenerationError as exc:
        detail = (
            _chatgpt_image_error_detail(exc)
            if display_model_key == ai_chat.CHATGPT_56_MODEL_KEY
            else str(exc)
        )
        response = JsonResponse(
            {'status': 'error', 'detail': detail},
            status=exc.status_code,
        )
        response['X-Conversation-Id'] = str(conversation.id)
        response['X-Model-Key'] = display_model_key
        response['X-Routed-Model-Key'] = _ai_public_routed_model_key(
            display_model_key, ai_chat.FLUX_KLEIN_4B_MODEL_KEY,
        )
        return response
    except Exception:
        logger.exception("Failed to save generated FLUX image")
        response = JsonResponse(
            {'status': 'error', 'detail': 'The image was generated but could not be saved. Please try again.'},
            status=503,
        )
        response['X-Conversation-Id'] = str(conversation.id)
        response['X-Model-Key'] = display_model_key
        response['X-Routed-Model-Key'] = _ai_public_routed_model_key(
            display_model_key, ai_chat.FLUX_KLEIN_4B_MODEL_KEY,
        )
        return response

    # No caption text — the image speaks for itself, and a canned "Image
    # generated successfully." heading above every single image was just
    # noise. The frontend (finalizeBubble in ai.html) knows not to fall
    # back to its usual "No response" placeholder when generatedImageUrl
    # is set, so an empty reply here renders as just the image.
    reply = ''
    assistant_message = AIMessage.objects.create(
        conversation=conversation,
        role=AIMessage.ROLE_ASSISTANT,
        content=reply,
        image_data=generated_url,
        model_key=display_model_key,
    )
    response = HttpResponse(reply, content_type='text/plain; charset=utf-8')
    response['Cache-Control'] = 'private, no-store'
    response['X-Conversation-Id'] = str(conversation.id)
    response['X-Model-Key'] = display_model_key
    response['X-Routed-Model-Key'] = _ai_public_routed_model_key(
        display_model_key, ai_chat.FLUX_KLEIN_4B_MODEL_KEY,
    )
    response['X-Request-Category'] = 'image_edit' if source_image else 'image_generation'
    response['X-Generated-Image-Url'] = generated_url
    response['X-Message-Id'] = str(assistant_message.pk)
    return response



def ai_chat_send(request):
    request_started = time.perf_counter()
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    # A real, billable API key sits behind this — a basic per-IP rate limit
    # is the brake against an account (or a guest, or a script cycling
    # through either) hammering it. Uses Django's cache, which on a
    # single-process dev server is exact; under multiple gunicorn workers
    # each worker has its own count, so the effective ceiling is
    # (limit × worker count) — a soft brake, not a hard guarantee.
    #
    # A fixed window with its own stored reset_at (rather than just an
    # integer whose cache TTL gets renewed on every message) so the exact
    # wait time can be told to whoever's blocked, instead of a vague "wait a
    # bit" — and so someone sending at a slow, steady trickle isn't kept
    # perpetually blocked by their own TTL renewing before it ever expires.
    ip = _client_ip(request)
    # Staff-issued blocks (see dashboard AI Activity) — checked before the
    # rate limiter below since a blocked spammer shouldn't even get to
    # accrue against it. Staff are never checked, so a mistaken block can
    # never accidentally lock out someone who can undo it.
    if not (request.user.is_authenticated and request.user.is_staff):
        block_filter = Q(ip_address=ip) if ip and ip != 'unknown' else Q(pk__isnull=True)
        if request.user.is_authenticated:
            block_filter |= Q(user=request.user)
        if AIBlock.objects.filter(block_filter).exists():
            return JsonResponse(
                {'status': 'error', 'detail': "Your access to Vidhyora AI has been restricted. Contact support if you think this is a mistake."},
                status=403,
            )
    cache_key = f'ai_chat_rate:{ip}'
    now = time.time()
    window = cache.get(cache_key)
    if not window or now >= window['reset_at']:
        window = {'count': 0, 'reset_at': now + AI_CHAT_RATE_WINDOW}
    if window['count'] >= AI_CHAT_RATE_LIMIT:
        wait_for = _format_wait_time(window['reset_at'] - now)
        return JsonResponse(
            {
                'status': 'rate_limited',
                'detail': f"Limit reached — please wait {wait_for} and try again later.",
            },
            status=429,
        )
    # Counted immediately, before any further validation — a request that
    # gets rejected downstream (bad payload, guest cap, etc.) still cost a
    # request and should still count against the brake, otherwise a blocked
    # guest can hit this endpoint an unlimited number of times for free.
    window['count'] += 1
    cache.set(cache_key, window, AI_CHAT_RATE_WINDOW)

    payload = _parse_json_body(request)
    if not isinstance(payload, dict):
        return JsonResponse({'status': 'error', 'detail': 'Invalid request body.'}, status=400)
    message = str(payload.get('message', ''))[:AI_CHAT_MAX_MESSAGE_CHARS].strip()

    image_data = payload.get('image')
    if not isinstance(image_data, str) or not image_data.startswith('data:image/'):
        image_data = ''
    if image_data and len(image_data) > AI_IMAGE_MAX_DATA_URI_CHARS:
        return JsonResponse({'status': 'error', 'detail': 'That image is too large — please use a smaller one.'}, status=400)

    # OCR supplements the multimodal model for screenshots and documents.
    # Failure is harmless: the original image is still analysed by Vision.
    image_ocr_text = image_ocr.extract_data_uri(image_data) if image_data else ''

    # document_text/document_name come from a prior call to /AI/api/extract/
    # (the raw file itself is never sent here) — re-capped defensively since
    # the client is untrusted, even though it already capped it once too.
    document_text = payload.get('document_text')
    document_name = payload.get('document_name')
    if isinstance(document_text, str) and document_text.strip() and isinstance(document_name, str) and document_name.strip():
        document_text = document_text.strip()[:doc_extract.MAX_CHARS]
        document_name = document_name.strip()[:255]
    else:
        document_text = ''
        document_name = ''
    requested_document_mode = payload.get('document_mode')
    document_mode = requested_document_mode if document_text and requested_document_mode in AI_DOCUMENT_MODES else ''
    document_truncated = bool(payload.get('document_truncated')) if document_text else False
    document_instruction = _ai_document_instruction(document_mode, document_name, document_truncated)

    if not message and not image_data and not document_text:
        return JsonResponse({'status': 'error', 'detail': 'No message provided.'}, status=400)

    requested_language = payload.get('language')
    language = requested_language if requested_language in ai_chat.LANGUAGES else ai_chat.DEFAULT_LANGUAGE

    # ChatGPT 5.6 is a stable user-facing selection backed by the existing
    # task-specific workers. Keep its public identity while routing the actual
    # turn to Vision, Code, or Quick.
    full_model_access = _ai_has_full_model_access(request.user)
    default_model_key = ai_chat.DEFAULT_MODEL_KEY if full_model_access else 'quick'
    requested_model_key = payload.get('model')
    selected_model_key = (
        requested_model_key
        if requested_model_key in ai_chat.MODELS
        else default_model_key
    )
    chatgpt_mode = selected_model_key == ai_chat.CHATGPT_56_MODEL_KEY
    response_model_key = ai_chat.CHATGPT_56_MODEL_KEY if chatgpt_mode else None

    # Gated on ai_chat.is_image_generation_request rather than just "FLUX is
    # selected" — that regex is what decides whether a message genuinely
    # reads as an image create/edit request (verb+noun within 40 chars,
    # bidirectional for Hindi/Hinglish word order) before anything gets
    # routed to an image-only model, so a plain "hi" doesn't get treated as
    # a picture prompt just because FLUX happened to be selected, and a real
    # "generate a poster for..." gets routed to FLUX even from a completely
    # different model without the user having to switch manually. A generic
    # capability question ("can you edit or generate images?") matches the
    # same regex but isn't an actual subject to draw, so it's excluded here
    # and answered conversationally instead — otherwise FLUX would try to
    # render the question text itself as an image.
    is_image_request = bool(message) and (
        (
            ai_chat.is_image_generation_request(message)
            and not ai_chat.is_image_capability_question(message)
        )
        # A raw Midjourney/DALL-E-style scene description (no generate/create
        # verb at all, e.g. "A realistic brown dog..., soft lighting, 4k.")
        # never matches the verb+noun regex above — see AIReport #19. Two or
        # more photography/art style cues and no question mark is a strong
        # enough signal to route it the same way.
        or (not image_data and ai_chat.is_probable_image_prompt(message))
    )
    generated_file_spec = (
        _ai_generated_file_spec(message)
        if message and not image_data and not document_text and not is_image_request
        else None
    )
    if generated_file_spec:
        document_instruction = _ai_generated_file_instruction(generated_file_spec['file_name'])

    if generated_file_spec:
        # A real server-created download must behave the same from every
        # picker choice (including FLUX, which cannot produce text files).
        model_key = 'code'
        request_category = 'file_generation'
    elif selected_model_key == ai_chat.FLUX_KLEIN_4B_MODEL_KEY:
        if image_data or (message and not ai_chat.is_image_capability_question(message)) or not message:
            # Once the user deliberately selects FLUX, descriptive prompts
            # such as "a robot in a futuristic classroom" are valid even
            # without an explicit generate/draw verb. Attachments are edits;
            # a bare capability question still falls back to normal chat.
            model_key = ai_chat.FLUX_KLEIN_4B_MODEL_KEY
            request_category = 'image_edit' if image_data else 'image_generation'
        else:
            # FLUX can only generate/edit images — it has no chat capability
            # at all, so a plain message typed while it happens to be
            # selected ("hi", a real question) is answered normally by
            # Quick instead of being forced through the image pipeline.
            model_key = 'quick'
            request_category = request_router.classify(message)
    elif is_image_request:
        # Any other model selected, but the message itself is genuinely
        # asking to create/edit an image — auto-route straight to FLUX
        # instead of just telling the user to switch models manually.
        model_key = ai_chat.FLUX_KLEIN_4B_MODEL_KEY
        request_category = 'image_edit' if image_data else 'image_generation'
    elif image_data:
        if ai_chat.is_image_edit_instruction(message):
            # An attached image already supplies the implicit subject, so an
            # edit-shaped verb alone is enough — "REmove all Names" on an
            # attached invitation (AIReport #11) has no image-shaped noun and
            # never matched is_image_request above, but sending it to Vision
            # (read-only) instead of FLUX (edit-capable) is exactly why it
            # got an unrelated reply instead of an edit.
            model_key = ai_chat.FLUX_KLEIN_4B_MODEL_KEY
            request_category = 'image_edit'
        else:
            model_key = 'vision'
    elif document_mode == 'coding':
        # "Start coding" is an explicit mode choice, so use the code-tuned
        # route even if the general model picker was previously on Light/etc.
        model_key = 'code'
        request_category = 'coding'
    else:
        model_key = selected_model_key
        if chatgpt_mode:
            model_key, request_category = request_router.choose_chatgpt_worker(message)
        elif model_key == ai_chat.DEFAULT_MODEL_KEY and payload.get('auto_route', True):
            model_key, request_category = request_router.choose_model(message, model_key)
        else:
            request_category = request_router.classify(message)

    if response_model_key is None:
        response_model_key = model_key

    if not full_model_access and (
        selected_model_key not in AI_FREE_MODEL_KEYS or model_key not in AI_FREE_MODEL_KEYS
    ):
        return JsonResponse({
            'status': 'subscription_required',
            'detail': (
                'The free plan includes Vidhyora Quick, Light, and Code. '
                'Premium access is required for this model or capability.'
            ),
        }, status=403)

    if model_key == ai_chat.FLUX_KLEIN_4B_MODEL_KEY and not message:
        return JsonResponse({
            'status': 'error',
            'detail': 'Describe the image you want, or how you want the attached image changed.',
        }, status=400)

    if not request.user.is_authenticated:
        if not request.session.session_key:
            request.session.create()
        # IP-derived rather than the session counter alone — the session
        # counter still gets tracked below for the on-page meter, but an
        # incognito window (or just clearing cookies) gets a fresh session
        # for free, so it can't be what actually gates access. Falls back to
        # the session counter only if the IP genuinely couldn't be read.
        guest_count = _ip_free_messages_used(ip) if ip and ip != 'unknown' else request.session.get('ai_guest_msg_count', 0)
        if guest_count >= AI_GUEST_MESSAGE_LIMIT:
            return JsonResponse({
                'status': 'login_required',
                'detail': "You've reached the free message limit — log in or sign up to keep chatting. Your conversation is saved and will carry over.",
            }, status=403)
    else:
        gate = _ai_profile_gate(request.user, ip)
        if gate:
            return JsonResponse(gate, status=403)

    owner_filter = _ai_owner_filter(request)
    conversation_id = payload.get('conversation_id')
    conversation = None
    if conversation_id:
        # payload is untrusted JSON — conversation_id could be a string, a
        # float, a list/dict, etc. Casting explicitly here means a malformed
        # value is just "not found" instead of an unhandled TypeError from
        # the ORM's pk lookup (which, with DEBUG on, would otherwise hand
        # back a full stack trace to whoever sent it).
        try:
            conversation_id = int(conversation_id)
        except (TypeError, ValueError):
            return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=404)
        conversation = AIConversation.objects.filter(owner_filter, pk=conversation_id).first()
        if not conversation:
            return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=404)
    if conversation is None:
        title = message[:AI_CONVERSATION_TITLE_CHARS] if message else (document_name or 'Image')
        conv_ip = ip if ip and ip != 'unknown' else None
        if request.user.is_authenticated:
            conversation = AIConversation.objects.create(user=request.user, title=title, ip_address=conv_ip)
        else:
            conversation = AIConversation.objects.create(session_key=request.session.session_key, title=title, ip_address=conv_ip)

    AIMessage.objects.create(
        conversation=conversation, role=AIMessage.ROLE_USER, content=message,
        image_data=image_data, document_name=document_name,
        document_text=document_text or image_ocr_text,
    )
    conversation.updated_at = timezone.now()
    conversation.save(update_fields=['updated_at'])
    request.session[AI_CURRENT_CONVERSATION_SESSION_KEY] = conversation.id

    # My Notes: 'show my notes' / 'delete note about X' / 'edit note about X
    # to Y' / 'take this note' — handled entirely here, no AI model call, so
    # none of it costs a free-message credit. Skipped whenever an image or
    # document is attached so that kind of turn can never get mistaken for
    # note-taking.
    if message and not image_data and not document_text:
        pending_response = _ai_pending_note_edit_response(request, conversation, message)
        if pending_response is not None:
            return pending_response
        if request_router.is_show_notes_intent(message):
            return _ai_show_notes_response(request, conversation)
        read_target = request_router.match_read_note(message)
        if read_target is not None:
            return _ai_read_note_response(request, conversation, read_target)
        delete_target = request_router.match_delete_note(message)
        if delete_target is not None:
            return _ai_delete_note_response(request, conversation, delete_target)
        edit_match = request_router.match_edit_note(message)
        if edit_match is not None:
            return _ai_edit_note_response(request, conversation, edit_match[0], edit_match[1], message)
        if request_router.is_contextual_note_edit(message):
            contextual_response = _ai_contextual_note_edit_response(request, conversation, message)
            if contextual_response is not None:
                return contextual_response
        if request_router.is_note_intent(message):
            return _ai_save_note_response(request, conversation, message)

    if not request.user.is_authenticated:
        request.session['ai_guest_msg_count'] = request.session.get('ai_guest_msg_count', 0) + 1
    elif not request.user.is_staff:
        # Only counts against the free-tier cap while unsubscribed — a
        # subscribed account's messages shouldn't erode the free allotment
        # that's waiting for them once the subscription lapses.
        StoreProfile.objects.filter(user=request.user).exclude(
            ai_subscription_until__gt=timezone.now(),
        ).update(ai_free_messages_used=F('ai_free_messages_used') + 1)

    if model_key == ai_chat.FLUX_KLEIN_4B_MODEL_KEY:
        image_gate_detail = _ai_image_gen_gate(request.user, ip)
        if image_gate_detail:
            response = JsonResponse({'status': 'rate_limited', 'detail': image_gate_detail}, status=429)
            response['X-Conversation-Id'] = str(conversation.id)
            response['X-Model-Key'] = response_model_key
            response['X-Routed-Model-Key'] = _ai_public_routed_model_key(
                response_model_key, ai_chat.FLUX_KLEIN_4B_MODEL_KEY,
            )
            return response
        return _ai_flux_response(conversation, message, image_data, response_model_key)

    # First-time-in-AI-chat onboarding: ask a genuinely new user (once) for
    # their name/location/Instagram, then deterministically capture whatever
    # they give in their very next reply — never trust the model itself to
    # record it, same reasoning as the My Notes system above. Skipped for
    # staff (site admins/developers, not real customers) and guests (there's
    # no account to save it against).
    onboarding_ask = False
    if request.user.is_authenticated and not request.user.is_staff:
        onboarding_profile, _ = StoreProfile.objects.get_or_create(user=request.user)
        if onboarding_profile.ai_onboarding_pending:
            fields = ai_chat.extract_onboarding_fields(message) if message else {}
            update_fields = ['ai_onboarding_pending', 'ai_onboarded']
            onboarding_profile.ai_onboarding_pending = False
            onboarding_profile.ai_onboarded = True
            if fields.get('name'):
                onboarding_profile.ai_display_name = fields['name'][:100]
                update_fields.append('ai_display_name')
            if fields.get('location'):
                onboarding_profile.ai_location = fields['location'][:150]
                update_fields.append('ai_location')
            if fields.get('instagram'):
                onboarding_profile.ai_instagram_handle = fields['instagram'].lstrip('@')[:60]
                update_fields.append('ai_instagram_handle')
            onboarding_profile.save(update_fields=update_fields)
        elif not onboarding_profile.ai_onboarded:
            onboarding_ask = True
            onboarding_profile.ai_onboarding_pending = True
            onboarding_profile.save(update_fields=['ai_onboarding_pending'])

    recent = list(
        conversation.messages.order_by('-created_at')
        .values('role', 'content', 'image_data', 'document_name', 'document_text')[:AI_CHAT_MAX_HISTORY]
    )
    recent.reverse()

    model_cfg = ai_chat.MODELS[model_key]

    clean_history = []
    for index, m in enumerate(recent):
        is_current_image = bool(m['image_data']) and index == len(recent) - 1
        if is_current_image and model_cfg['vision']:
            ocr_note = f"\n\nLocally detected text:\n{m['document_text']}" if m['document_text'] else ''
            content = [
                {'type': 'text', 'text': (m['content'] or 'Describe and analyse this image.') + ocr_note},
                {'type': 'image_url', 'image_url': {'url': m['image_data']}},
            ]
        elif m['image_data'] and m['role'] == AIMessage.ROLE_ASSISTANT:
            # This is a generated media URL, not a new user attachment.
            content = (m['content'] + ' [An image was generated in this turn.]').strip()
        elif m['image_data']:
            # Historical base64 image bytes are not resent. The earlier
            # assistant analysis remains in history, with OCR as backup.
            ocr_note = f" Earlier image OCR:\n{m['document_text']}" if m['document_text'] else ''
            content = (m['content'] + ' [Earlier image was analysed.]' + ocr_note).strip()
        elif m['document_name'] and m['document_text']:
            # Persisted in full (capped by doc_extract.MAX_CHARS) so a
            # follow-up question about this document works without
            # re-uploading it, for as long as it stays within the replayed
            # AI_CHAT_MAX_HISTORY window.
            content = (
                f"[Attached document: {m['document_name']}]\n{m['document_text']}\n\n---\n"
                f"{m['content'] or 'Please review the attached document.'}"
            )
        elif m['document_name']:
            # Older row from before document_text was persisted — degrade
            # to a plain filename note instead of silently having nothing.
            content = (m['content'] + f" [Attached document: {m['document_name']}]").strip()
        else:
            content = m['content']
        clean_history.append({'role': m['role'], 'content': content})

    # Drop the oldest turns (documents/images first pushed the total over
    # budget) until the whole history fits, always keeping at least the
    # current turn — a long-running conversation with attachments degrades
    # gracefully instead of eventually overflowing the model's context
    # window (see AI_CHAT_HISTORY_CHAR_BUDGET above).
    def _content_char_len(content):
        if isinstance(content, str):
            return len(content)
        if isinstance(content, list):
            return sum(len(b.get('text', '')) for b in content if isinstance(b, dict))
        return 0

    while len(clean_history) > 1 and (
        sum(_content_char_len(m['content']) for m in clean_history) > AI_CHAT_HISTORY_CHAR_BUDGET
    ):
        clean_history.pop(0)

    # Keep saved conversation text intact for the user, but redact common
    # personal identifiers in the copy sent to the external model.
    for history_message in clean_history:
        content = history_message['content']
        if isinstance(content, str):
            history_message['content'] = privacy.redact(content)
        elif isinstance(content, list):
            for block in content:
                if block.get('type') == 'text':
                    block['text'] = privacy.redact(block.get('text', ''))

    user_context = _ai_user_context(request.user) if request.user.is_authenticated else None
    if user_context:
        user_context = privacy.redact(user_context)

    # Once the secret phrase has appeared anywhere in this conversation, the
    # warm/personal tone stays on for the rest of it rather than resetting
    # every message. The greeting/introduction itself, though, should only
    # happen on the message where she actually said the phrase — not be
    # repeated on every later message just because the tone stays on.
    #
    # Sumudrika and Jagu are deliberately mutually exclusive within one
    # conversation — whichever trigger phrase appears EARLIEST wins and
    # locks in for the rest of it, so a later "my name is jagu" mentioned
    # mid-Sumudrika-conversation (or vice versa) can't make the model try
    # to address two different people as the same "her" and get confused.
    sumudrika_idx = next(
        (i for i, m in enumerate(recent) if m['role'] == AIMessage.ROLE_USER and ai_chat.is_sumudrika_trigger(m['content'])),
        None,
    )
    jagu_idx = next(
        (i for i, m in enumerate(recent) if m['role'] == AIMessage.ROLE_USER and ai_chat.is_jagu_trigger(m['content'])),
        None,
    )
    if sumudrika_idx is not None and (jagu_idx is None or sumudrika_idx <= jagu_idx):
        is_sumudrika, is_jagu = True, False
    elif jagu_idx is not None:
        is_sumudrika, is_jagu = False, True
    else:
        is_sumudrika, is_jagu = False, False

    is_sumudrika_greet = is_sumudrika and recent[-1]['role'] == AIMessage.ROLE_USER \
        and ai_chat.is_sumudrika_trigger(recent[-1]['content'])
    is_jagu_greet = is_jagu and recent[-1]['role'] == AIMessage.ROLE_USER \
        and ai_chat.is_jagu_trigger(recent[-1]['content'])
    # Checked only against the current message, not the whole history — see
    # ai_chat.is_persona_farewell.
    is_persona_farewell = (is_sumudrika or is_jagu) and bool(recent) \
        and recent[-1]['role'] == AIMessage.ROLE_USER and ai_chat.is_persona_farewell(recent[-1]['content'])

    # The Sumudrika/Jagu personas depend on subtle instruction-following
    # (warm tone, correct pronouns, never fabricating a quote from Rudra)
    # that live-testing showed the smaller/faster modes don't reliably
    # deliver — EduTrellis Quick (the site default) was caught inventing
    # fake quotes attributed to Rudra and misgendering him. These are rare,
    # low-volume, personally-important conversations, so correctness wins
    # over speed/cost here: force EduTrellis Ultra once triggered,
    # regardless of whatever mode was actually selected — except when an
    # image is attached this turn, since Ultra has no vision capability and
    # that has to win.
    if full_model_access and (is_sumudrika or is_jagu) and model_key != 'vision':
        model_key = 'ultra'

    # Verified, static business-fact grounding only — never a saved/editable
    # answer store. No AI reply is ever generated from a knowledge base or
    # cached past answer.
    retrieved_context = retrieved_source = None
    recent_company_text = ' '.join(
        str(item.get('content') or '') for item in recent[-5:]
        if isinstance(item.get('content'), str)
    )
    if message and company_knowledge.is_company_query(recent_company_text):
        retrieved_context = company_knowledge.PUBLIC_SITE_CONTEXT
        retrieved_source = 'company_site'

    logger.info(
        "AI timing preprocessing=%.3fs model=%s history=%d image=%s ocr_chars=%d",
        time.perf_counter() - request_started, model_key, len(recent),
        bool(image_data), len(image_ocr_text),
    )

    def event_stream():
        full_reply = ''
        had_error = False
        hide_chatgpt_worker = response_model_key == ai_chat.CHATGPT_56_MODEL_KEY
        try:
            for chunk in ai_chat.stream_chat(
                clean_history, model_key=model_key,
                identity_model_key=(response_model_key if response_model_key != model_key else None),
                user_context=user_context,
                retrieved_context=retrieved_context, retrieved_source=retrieved_source,
                sumudrika=is_sumudrika, sumudrika_greet=is_sumudrika_greet,
                jagu=is_jagu, jagu_greet=is_jagu_greet,
                persona_farewell=is_persona_farewell, language=language,
                document_instruction=document_instruction,
                max_tokens=(
                    AI_DOCUMENT_CODE_MAX_OUTPUT_TOKENS
                    if document_mode == 'coding' or generated_file_spec or ai_chat.wants_long_form_output(message)
                    else None
                ),
                onboarding_ask=onboarding_ask,
            ):
                full_reply += chunk
                if not hide_chatgpt_worker:
                    yield chunk
            if hide_chatgpt_worker:
                # Buffer this one public identity until the upstream reply is
                # complete. That lets us catch a hidden model name even when
                # it is split across streaming chunks (for example "FL" +
                # "UX") before any of it reaches the browser.
                full_reply = _chatgpt_public_reply(full_reply)
                if full_reply:
                    yield full_reply
            if generated_file_spec and full_reply.strip():
                try:
                    generated_file = AIGeneratedFile.objects.create(
                        user=request.user if request.user.is_authenticated else None,
                        session_key='' if request.user.is_authenticated else (request.session.session_key or ''),
                        file_name=generated_file_spec['file_name'],
                        content=_extract_ai_generated_file_content(full_reply),
                    )
                    download_url = request.build_absolute_uri(reverse(
                        'ai_generated_file_download', args=[generated_file.token],
                    ))
                    download_link = f"\n\n[Download {generated_file.file_name}]({download_url})"
                    full_reply += download_link
                    yield download_link
                except Exception:
                    # The model response is still useful if persistence ever
                    # fails; don't mislabel a completed answer as a stream
                    # failure merely because its download could not be saved.
                    logger.exception("Failed to save an AI-generated download")
        except Exception as e:
            # ai_chat.stream_chat already retries transient failures on its
            # own before ever raising here — reaching this point means every
            # retry failed (or a real answer had already started streaming
            # when it broke, so a clean retry wasn't possible). Either way,
            # what's in full_reply (if anything) is not a complete, trustworthy
            # answer, so it must never be saved as one.
            logger.exception("AI chat stream failed after retries: %s", e)
            had_error = True
            if full_reply.strip():
                # A dropped mobile connection or upstream stream can happen
                # after useful text has arrived. Keep that text visible
                # instead of replacing it with a generic failure message.
                if hide_chatgpt_worker:
                    yield _chatgpt_public_reply(full_reply)
                yield "\n\n[Response interrupted. You can retry if anything is missing.]"
            elif ai_chat._is_context_length_error(e):
                # Retrying would just fail again identically — the fixed
                # generic message was misleading here, since it reads as
                # transient when the real fix is a shorter message/thread.
                yield (
                    "This conversation (or an attached document) has gotten "
                    "too long for the AI to process in one go. Please start "
                    "a new chat, or ask about a shorter excerpt."
                )
            else:
                yield _ai_chat_failure_reply(
                    e, response_model_key,
                    is_staff=bool(request.user.is_authenticated and request.user.is_staff),
                )
        finally:
            # The conversation can have been deleted (by this same user, in
            # another tab, or via the sidebar delete button) while this reply
            # was still streaming — check it still exists before trying to
            # attach a message to it, instead of letting that blow up here.
            if full_reply.strip() and not had_error:
                try:
                    if AIConversation.objects.filter(pk=conversation.pk).exists():
                        AIMessage.objects.create(
                            conversation=conversation, role=AIMessage.ROLE_ASSISTANT,
                            content=full_reply, model_key=response_model_key,
                        )
                except Exception:
                    logger.exception("Failed to save AI assistant reply for conversation %s", conversation.pk)

    response = StreamingHttpResponse(event_stream(), content_type='text/plain; charset=utf-8')
    response['Cache-Control'] = 'no-cache'
    response['X-Accel-Buffering'] = 'no'
    response['X-Conversation-Id'] = str(conversation.id)
    response['X-Model-Key'] = response_model_key
    response['X-Routed-Model-Key'] = _ai_public_routed_model_key(response_model_key, model_key)
    response['X-Request-Category'] = request_category if not image_data else 'image'
    # Tells the frontend to auto-play this reply and show the persona
    # follow-up chips — true for every turn once the matching trigger
    # phrase has appeared anywhere in the conversation (same scope as
    # is_sumudrika/is_jagu above, not just the one turn that said it).
    response['X-Sumudrika'] = '1' if is_sumudrika else ''
    response['X-Jagu'] = '1' if is_jagu else ''
    # Tells the frontend this was her goodbye reply — lock the composer
    # instead of showing the usual follow-up chips (see is_persona_farewell
    # above; ai_chat.stream_chat was told to make this a closing message).
    response['X-Persona-End'] = '1' if is_persona_farewell else ''
    return response


AI_DOC_RATE_LIMIT = 15
AI_DOC_MAX_UPLOAD_BYTES = 8 * 1024 * 1024  # raw file cap, before extraction


def ai_extract_document(request):
    """Extracts text from an uploaded PDF/DOCX/CSV and hands it back to the
    browser — the raw file is never stored or forwarded anywhere else, and
    the caller sends the returned text back in the next /AI/api/send/ call
    (see ai_chat_send's document_text/document_name handling)."""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    ip = _client_ip(request)
    cache_key = f'ai_doc_rate:{ip}'
    count = cache.get(cache_key, 0)
    if count >= AI_DOC_RATE_LIMIT:
        return JsonResponse(
            {'status': 'error', 'detail': "You're uploading files too quickly — please wait a bit and try again."},
            status=429,
        )
    cache.set(cache_key, count + 1, AI_CHAT_RATE_WINDOW)

    f = request.FILES.get('file')
    if not f:
        return JsonResponse({'status': 'error', 'detail': 'No file provided.'}, status=400)
    if f.size > AI_DOC_MAX_UPLOAD_BYTES:
        return JsonResponse({'status': 'error', 'detail': 'That file is too large — please use one under 8MB.'}, status=400)

    try:
        file_bytes = f.read()
        text, truncated = doc_extract.extract(f.name, file_bytes)
        coding_text, coding_truncated = doc_extract.extract_editable_source(
            f.name, file_bytes, text, extracted_truncated=truncated,
        )
    except doc_extract.ExtractError as e:
        return JsonResponse({'status': 'error', 'detail': str(e)}, status=400)
    except Exception:
        logger.exception("Document extraction failed for %s", f.name)
        return JsonResponse({'status': 'error', 'detail': "Couldn't read that file — please try a different one."}, status=400)

    return JsonResponse({
        'status': 'ok', 'filename': f.name,
        'text': text, 'truncated': truncated,
        'coding_text': coding_text, 'coding_truncated': coding_truncated,
    })


def ai_transcribe_audio(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    uploaded = request.FILES.get('file')
    if not uploaded:
        return JsonResponse({'status': 'error', 'detail': 'No audio file provided.'}, status=400)
    if uploaded.size > 12 * 1024 * 1024:
        return JsonResponse({'status': 'error', 'detail': 'Audio must be under 12MB.'}, status=400)
    try:
        text = audio_transcribe.transcribe(uploaded.name, uploaded.read())
    except RuntimeError as exc:
        return JsonResponse({'status': 'error', 'detail': str(exc)}, status=400)
    if not text:
        return JsonResponse({'status': 'error', 'detail': 'No speech was detected.'}, status=400)
    return JsonResponse({'status': 'ok', 'text': text})


def _cleanup_youtube_downloads():
    expired = list(YouTubeDownloadJob.objects.filter(expires_at__lt=timezone.now()))
    media_root = Path(settings.MEDIA_ROOT).resolve()
    for job in expired:
        for relative in (job.video_path, job.audio_path):
            if not relative:
                continue
            path = (media_root / relative).resolve()
            if media_root in path.parents:
                try:
                    path.unlink(missing_ok=True)
                except OSError:
                    pass
        job.delete()


def ai_youtube_start(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'login_required', 'detail': 'Log in to prepare downloads.'}, status=403)
    payload = _parse_json_body(request)
    url = str(payload.get('url', '')).strip() if isinstance(payload, dict) else ''
    quality = str(payload.get('quality', '1080')).strip() if isinstance(payload, dict) else '1080'
    if not youtube_download.is_youtube_url(url):
        return JsonResponse({'status': 'error', 'detail': 'Enter a valid YouTube video URL.'}, status=400)
    if quality not in ('720', '1080', 'audio'):
        return JsonResponse({'status': 'error', 'detail': 'Choose 720p, 1080p, or MP3.'}, status=400)
    ip = _client_ip(request)
    rate_key = f'ai_youtube_rate:{request.user.pk}:{ip}'
    count = cache.get(rate_key, 0)
    if count >= 3:
        return JsonResponse({'status': 'error', 'detail': 'Download limit reached — try again in one hour.'}, status=429)
    if YouTubeDownloadJob.objects.filter(user=request.user, status__in=['pending', 'working']).exists():
        return JsonResponse({'status': 'error', 'detail': 'Another video is already being prepared.'}, status=409)
    _cleanup_youtube_downloads()
    job = YouTubeDownloadJob.objects.create(
        user=request.user, source_url=url, quality=quality,
        expires_at=timezone.now() + timedelta(hours=1),
    )
    cache.set(rate_key, count + 1, 60 * 60)
    youtube_download.start(job.pk)
    return JsonResponse({'status': 'ok', 'job': str(job.token)}, status=202)


def _youtube_job_for_user(request, token):
    if not request.user.is_authenticated:
        return None
    return YouTubeDownloadJob.objects.filter(user=request.user, token=token).first()


def ai_youtube_status(request, token):
    job = _youtube_job_for_user(request, token)
    if not job:
        return JsonResponse({'status': 'error', 'detail': 'Download not found.'}, status=404)
    payload = {'status': job.status, 'progress': job.progress, 'title': job.title, 'error': job.error}
    if job.status == YouTubeDownloadJob.STATUS_READY:
        if job.video_path:
            payload['video_url'] = f'/AI/api/youtube/{job.token}/file/video/'
        if job.audio_path:
            payload['audio_url'] = f'/AI/api/youtube/{job.token}/file/audio/'
        payload['expires_at'] = timezone.localtime(job.expires_at).isoformat()
    return JsonResponse(payload)


def ai_youtube_file(request, token, file_kind):
    job = _youtube_job_for_user(request, token)
    if not job or job.status != YouTubeDownloadJob.STATUS_READY or job.expires_at <= timezone.now():
        return JsonResponse({'status': 'error', 'detail': 'Download unavailable or expired.'}, status=404)
    relative = job.video_path if file_kind == 'video' else job.audio_path if file_kind == 'audio' else ''
    media_root = Path(settings.MEDIA_ROOT).resolve()
    path = (media_root / relative).resolve() if relative else None
    if not path or media_root not in path.parents or not path.is_file():
        return JsonResponse({'status': 'error', 'detail': 'File not found.'}, status=404)
    extension = path.suffix.lower()
    content_type = 'audio/mpeg' if file_kind == 'audio' else (mimetypes.guess_type(path.name)[0] or 'application/octet-stream')
    filename = f"{job.title or 'youtube-download'}{extension}"
    response = FileResponse(open(path, 'rb'), content_type=content_type, as_attachment=True, filename=filename)
    response['Cache-Control'] = 'private, no-store'
    return response


def ai_conversations_list(request):
    conversations = list(
        AIConversation.objects.filter(_ai_owner_filter(request))
        .values('id', 'title', 'updated_at').order_by('-updated_at')[:100]
    )
    for c in conversations:
        c['updated_at'] = timezone.localtime(c['updated_at']).isoformat()
    return JsonResponse({'status': 'ok', 'conversations': conversations})


def ai_conversation_messages(request, conversation_id):
    conversation = AIConversation.objects.filter(_ai_owner_filter(request), pk=conversation_id).first()
    if not conversation:
        return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=404)
    request.session[AI_CURRENT_CONVERSATION_SESSION_KEY] = conversation.id
    messages_qs = list(
        conversation.messages.order_by('created_at')
        .values('id', 'role', 'content', 'image_data', 'document_name', 'model_key')
    )
    return JsonResponse({'status': 'ok', 'title': conversation.title, 'messages': messages_qs})


def ai_conversation_delete(request, conversation_id):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    conversation = AIConversation.objects.filter(_ai_owner_filter(request), pk=conversation_id).first()
    if not conversation:
        return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=404)
    conversation.delete()
    if request.session.get(AI_CURRENT_CONVERSATION_SESSION_KEY) == conversation_id:
        request.session.pop(AI_CURRENT_CONVERSATION_SESSION_KEY, None)
    return JsonResponse({'status': 'ok'})


AI_IMPORT_MAX_MESSAGES = 500
AI_IMPORT_MAX_MESSAGE_CHARS = 8000


def ai_conversation_import(request):
    """Creates a new conversation from a chat transcript the user uploads
    from the account dropdown's "Import chat" option (see the Vidhyora AI
    account menu in ai.html) — a JSON body shaped either as a bare list of
    messages or {title, messages}, each message {role, content} with role
    'user'/'assistant' (a handful of common aliases are accepted client-side
    in normalizeImportedChat before this ever gets called). Owned the same
    dual way as a conversation started from chat itself, so an imported
    guest chat still carries over on login like any other."""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    payload = _parse_json_body(request)
    if not isinstance(payload, dict):
        return JsonResponse({'status': 'error', 'detail': 'Invalid import file.'}, status=400)
    raw_messages = payload.get('messages')
    if not isinstance(raw_messages, list) or not raw_messages:
        return JsonResponse({'status': 'error', 'detail': 'That file has no recognizable chat messages.'}, status=400)

    cleaned = []
    for item in raw_messages[:AI_IMPORT_MAX_MESSAGES]:
        if not isinstance(item, dict):
            continue
        role = item.get('role')
        content = item.get('content')
        if role not in (AIMessage.ROLE_USER, AIMessage.ROLE_ASSISTANT):
            continue
        if not isinstance(content, str) or not content.strip():
            continue
        cleaned.append((role, content[:AI_IMPORT_MAX_MESSAGE_CHARS]))
    if not cleaned:
        return JsonResponse({'status': 'error', 'detail': 'That file has no recognizable chat messages.'}, status=400)

    title = payload.get('title')
    title = title.strip()[:AI_CONVERSATION_TITLE_CHARS] if isinstance(title, str) and title.strip() else ''
    if not title:
        first_user = next((content for role, content in cleaned if role == AIMessage.ROLE_USER), cleaned[0][1])
        title = first_user[:AI_CONVERSATION_TITLE_CHARS]

    if not request.user.is_authenticated and not request.session.session_key:
        request.session.create()
    ip = _client_ip(request)
    conv_ip = ip if ip and ip != 'unknown' else None
    if request.user.is_authenticated:
        conversation = AIConversation.objects.create(user=request.user, title=title, ip_address=conv_ip)
    else:
        conversation = AIConversation.objects.create(session_key=request.session.session_key, title=title, ip_address=conv_ip)

    AIMessage.objects.bulk_create([
        AIMessage(conversation=conversation, role=role, content=content) for role, content in cleaned
    ])

    return JsonResponse({'status': 'ok', 'conversation': {'id': conversation.id, 'title': conversation.title}})


def ai_notes_list(request):
    response = JsonResponse({'status': 'ok', 'notes': _ai_notes_snapshot(request)})
    response['Cache-Control'] = 'private, no-store'
    return response


def ai_note_delete(request, note_id):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    note = AINote.objects.filter(_ai_owner_filter(request), pk=note_id).first()
    if not note:
        return JsonResponse({'status': 'error', 'detail': 'Note not found.'}, status=404)
    note.delete()
    if any(item['id'] == note_id for item in _ai_notes_snapshot(request)):
        return JsonResponse({'status': 'error', 'detail': 'Could not delete note.'}, status=500)
    response = JsonResponse({'status': 'ok', 'notes': _ai_notes_snapshot(request)})
    response['Cache-Control'] = 'private, no-store'
    return response


AI_REPORT_MAX_EXPLANATION_CHARS = 2000


def _is_reportable_unsaved_ai_failure(reply):
    """Allow reports for streamed provider failures that are not AIMessage rows."""
    lowered = (reply or '').casefold()
    return any(marker in lowered for marker in (
        'did not respond',
        'taking longer than expected',
        'temporarily unavailable',
        'currently disconnected',
        'authentication is currently unavailable',
        'currently at its request limit',
        '[response interrupted.',
        'too long for the ai to process',
    ))


def _ai_storage_name_from_url(image_value):
    """Return a safe storage-relative name for one of our media URLs."""
    if not image_value or image_value.startswith('data:'):
        return ''
    image_path = unquote(urlsplit(image_value).path or '')
    media_path = urlsplit(settings.MEDIA_URL or '/media/').path or '/media/'
    if not media_path.endswith('/'):
        media_path += '/'
    if not image_path.startswith(media_path):
        return ''
    storage_name = image_path[len(media_path):].lstrip('/')
    storage_path = Path(storage_name)
    if not storage_name or '\\' in storage_name or storage_path.is_absolute() or '..' in storage_path.parts:
        return ''
    return storage_name


def _snapshot_ai_report_image(image_value):
    """Put visual report evidence in the DB when its media file is readable."""
    if not image_value:
        return ''
    if image_value.startswith('data:image/'):
        return image_value[:AI_IMAGE_MAX_DATA_URI_CHARS]
    storage_name = _ai_storage_name_from_url(image_value)
    if not storage_name:
        return image_value
    try:
        with default_storage.open(storage_name, 'rb') as image_file:
            raw = image_file.read(AI_IMAGE_MAX_DATA_URI_CHARS + 1)
        if len(raw) > AI_IMAGE_MAX_DATA_URI_CHARS:
            return image_value
        content_type = mimetypes.guess_type(storage_name)[0] or 'image/png'
        return f'data:{content_type};base64,{base64.b64encode(raw).decode("ascii")}'
    except Exception:
        # The report is still useful if an old deployment already lost the
        # media file; retain its original URL and surface that health problem
        # to staff on the reports dashboard.
        logger.warning('Could not snapshot reported AI image %s', storage_name)
        return image_value


def ai_report_submit(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    payload = _parse_json_body(request)
    if not isinstance(payload, dict):
        return JsonResponse({'status': 'error', 'detail': 'Invalid request body.'}, status=400)

    conversation_id = payload.get('conversation_id')
    try:
        conversation_id = int(conversation_id)
    except (TypeError, ValueError):
        return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=400)
    conversation = AIConversation.objects.filter(_ai_owner_filter(request), pk=conversation_id).first()
    if not conversation:
        return JsonResponse({'status': 'error', 'detail': 'Conversation not found.'}, status=404)

    reported_reply = str(payload.get('reply_text', ''))[:AI_CHAT_MAX_MESSAGE_CHARS].strip()
    reported_image = payload.get('reply_image', '')
    if not isinstance(reported_image, str):
        reported_image = ''
    reported_image = reported_image[:AI_IMAGE_MAX_DATA_URI_CHARS].strip()
    explanation = str(payload.get('explanation', ''))[:AI_REPORT_MAX_EXPLANATION_CHARS].strip()
    model_key = str(payload.get('model_key', ''))[:20]
    if not reported_reply and not reported_image:
        return JsonResponse({'status': 'error', 'detail': 'Nothing to report.'}, status=400)
    if not explanation:
        return JsonResponse({'status': 'error', 'detail': 'Please explain what went wrong before submitting.'}, status=400)

    # Best-effort link to the actual saved AIMessage row, purely so staff can
    # jump straight to it from admin — the report is still saved (with its
    # own snapshot of the text) even when this doesn't find one, e.g. the
    # message has since been edited or the conversation deleted.
    raw_message_id = payload.get('message_id')
    message = None
    if raw_message_id not in (None, ''):
        try:
            message_id = int(raw_message_id)
        except (TypeError, ValueError):
            return JsonResponse({'status': 'error', 'detail': 'Reported response not found.'}, status=400)
        message = conversation.messages.filter(
            pk=message_id, role=AIMessage.ROLE_ASSISTANT,
        ).first()
        if message is None:
            return JsonResponse({'status': 'error', 'detail': 'Reported response not found.'}, status=404)
    else:
        candidates = conversation.messages.filter(role=AIMessage.ROLE_ASSISTANT)
        if reported_reply:
            candidates = candidates.filter(content=reported_reply)
        if reported_image:
            candidates = candidates.filter(image_data=reported_image)
        message = candidates.order_by('-created_at', '-pk').first()

    # Stored messages are authoritative: a browser cannot substitute a
    # different reply, model, or image in the report payload.
    if message is not None:
        reported_reply = message.content[:AI_CHAT_MAX_MESSAGE_CHARS].strip()
        reported_image = message.image_data
        model_key = message.model_key[:20]
    elif reported_image:
        # Image responses are persisted before being returned. An unmatched
        # URL/data URI therefore cannot be genuine evidence from this chat.
        return JsonResponse({'status': 'error', 'detail': 'Reported image not found.'}, status=404)
    elif not _is_reportable_unsaved_ai_failure(reported_reply):
        return JsonResponse({'status': 'error', 'detail': 'Reported response not found.'}, status=404)

    # Snapshot the user turn that prompted this reply, so staff reviewing
    # the report don't have to go dig through the (possibly since-deleted)
    # conversation to see what was actually asked. Anchored to the matched
    # message's timestamp when we have one; otherwise best-effort fall back
    # to the most recent user turn in the conversation.
    preceding = conversation.messages.filter(role=AIMessage.ROLE_USER)
    if message:
        preceding = preceding.filter(
            Q(created_at__lt=message.created_at) |
            Q(created_at=message.created_at, pk__lt=message.pk)
        )
    user_message = preceding.order_by('-created_at', '-pk').first()
    user_prompt = user_message.content[:AI_CHAT_MAX_MESSAGE_CHARS].strip() if user_message else ''
    user_image = _snapshot_ai_report_image(user_message.image_data) if user_message else ''
    user_document_name = user_message.document_name if user_message else ''
    user_document_excerpt = user_message.document_text[:8000] if user_message else ''
    reported_image = _snapshot_ai_report_image(reported_image)

    ip = _client_ip(request)
    AIReport.objects.create(
        conversation=conversation, message=message,
        user_prompt=user_prompt, user_image=user_image,
        user_document_name=user_document_name,
        user_document_excerpt=user_document_excerpt,
        reported_reply=reported_reply, reported_image=reported_image,
        model_key=model_key, explanation=explanation,
        user=request.user if request.user.is_authenticated else None,
        session_key='' if request.user.is_authenticated else (request.session.session_key or ''),
        ip_address=ip if ip and ip != 'unknown' else None,
    )
    return JsonResponse({'status': 'ok'})


# ── AI GitHub mode: connect a repo, let the AI CRUD + push to it on prompt ──
# Available to every authenticated account. Each user has a separate
# connection/token and all queries remain scoped to request.user, so no
# repository credentials or selections are shared between accounts.

AI_GITHUB_RATE_LIMIT = 10
AI_GITHUB_MAX_FILE_CHARS = 60_000  # skip pulling absurdly large files into the prompt


def _github_guard(request):
    return request.user.is_authenticated


def _github_forbidden():
    return JsonResponse({'status': 'error', 'detail': 'Log in to use GitHub.'}, status=401)


def github_status(request):
    if not _github_guard(request):
        return _github_forbidden()
    conn = GitHubConnection.objects.filter(user=request.user).first()
    if not conn:
        return JsonResponse({'status': 'ok', 'connected': False})
    return JsonResponse({
        'status': 'ok', 'connected': True, 'github_username': conn.github_username,
        'repo': conn.repo_full_name, 'branch': conn.default_branch,
    })


def _github_oauth_redirect_uri(request):
    return request.build_absolute_uri('/AI/api/github/oauth/callback/')


def _save_github_connection(user, token, github_username, repos):
    """Save credentials and keep only a repository visible to this token."""
    conn, _created = GitHubConnection.objects.update_or_create(
        user=user,
        defaults={'access_token': token, 'github_username': github_username},
    )
    available = {repo['full_name']: repo for repo in repos}
    selected = available.get(conn.repo_full_name)
    if selected is None:
        selected = next(
            (repo for repo in repos if 'edutrellis' in repo['full_name'].lower()),
            repos[0] if repos else None,
        )
    conn.repo_full_name = selected['full_name'] if selected else ''
    conn.default_branch = selected['default_branch'] if selected else 'main'
    conn.save(update_fields=['repo_full_name', 'default_branch'])
    return conn


def github_oauth_start(request):
    # A redirect, not a JSON endpoint — the browser navigates here directly
    # (window.location.href), so an unauthorized visit bounces to the
    # AI page instead of showing a raw 403 JSON body.
    if not _github_guard(request):
        return redirect('ai_page')
    if not settings.GITHUB_OAUTH_CLIENT_ID:
        return redirect('/AI/?github_error=not_configured')
    state = secrets.token_urlsafe(24)
    request.session['github_oauth_state'] = state
    params = {
        'client_id': settings.GITHUB_OAUTH_CLIENT_ID,
        'redirect_uri': _github_oauth_redirect_uri(request),
        # 'repo' grants full read/write on both private and public repos —
        # the closest standard GitHub OAuth scope to "all permission on the
        # repo", without also requesting unrelated org/user-admin scopes.
        'scope': 'repo',
        'state': state,
    }
    return redirect('https://github.com/login/oauth/authorize?' + urlencode(params))


def github_oauth_callback(request):
    if not _github_guard(request):
        return redirect('ai_page')
    expected_state = request.session.pop('github_oauth_state', None)
    state = request.GET.get('state')
    if not state or not expected_state or not secrets.compare_digest(state, expected_state):
        return redirect('/AI/?github_error=state')
    code = request.GET.get('code')
    if not code:
        return redirect('/AI/?github_error=denied')
    try:
        token_resp = requests.post(
            'https://github.com/login/oauth/access_token',
            data={
                'client_id': settings.GITHUB_OAUTH_CLIENT_ID,
                'client_secret': settings.GITHUB_OAUTH_CLIENT_SECRET,
                'code': code,
                'redirect_uri': _github_oauth_redirect_uri(request),
            },
            headers={'Accept': 'application/json'},
            timeout=20,
        )
        token = token_resp.json().get('access_token')
        if not token:
            return redirect('/AI/?github_error=token')
        gh_user = github_ops.get_authenticated_user(token)
        repos = github_ops.list_user_repos(token)
    except Exception:
        logger.exception("GitHub OAuth callback failed")
        return redirect('/AI/?github_error=1')

    _save_github_connection(request.user, token, gh_user.get('login', ''), repos)
    return redirect('/AI/?github_connected=1')


def github_connect(request):
    if not _github_guard(request):
        return _github_forbidden()
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    payload = _parse_json_body(request)
    if not isinstance(payload, dict):
        return JsonResponse({'status': 'error', 'detail': 'Invalid request body.'}, status=400)
    token = str(payload.get('token', '')).strip()
    if not token:
        return JsonResponse({'status': 'error', 'detail': 'Token required.'}, status=400)
    try:
        gh_user = github_ops.get_authenticated_user(token)
        repos = github_ops.list_user_repos(token)
    except github_ops.GitHubAPIError as e:
        return JsonResponse({'status': 'error', 'detail': f"Couldn't connect: {e}"}, status=400)
    conn = _save_github_connection(request.user, token, gh_user.get('login', ''), repos)
    return JsonResponse({
        'status': 'ok', 'github_username': conn.github_username, 'repos': repos,
        'repo': conn.repo_full_name, 'branch': conn.default_branch,
    })


def github_repos(request):
    if not _github_guard(request):
        return _github_forbidden()
    conn = GitHubConnection.objects.filter(user=request.user).first()
    if not conn:
        return JsonResponse({'status': 'error', 'detail': 'Not connected.'}, status=400)
    try:
        repos = github_ops.list_user_repos(conn.access_token)
    except github_ops.GitHubAPIError as e:
        return JsonResponse({'status': 'error', 'detail': str(e)}, status=400)
    return JsonResponse({'status': 'ok', 'repos': repos})


def github_set_repo(request):
    if not _github_guard(request):
        return _github_forbidden()
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    conn = GitHubConnection.objects.filter(user=request.user).first()
    if not conn:
        return JsonResponse({'status': 'error', 'detail': 'Not connected.'}, status=400)
    payload = _parse_json_body(request)
    repo = str(payload.get('repo', '')).strip() if isinstance(payload, dict) else ''
    if '/' not in repo:
        return JsonResponse({'status': 'error', 'detail': 'Invalid repo.'}, status=400)
    owner, _, name = repo.partition('/')
    try:
        info = github_ops.get_repo(conn.access_token, owner, name)
    except github_ops.GitHubAPIError as e:
        return JsonResponse({'status': 'error', 'detail': str(e)}, status=400)
    conn.repo_full_name = info['full_name']
    conn.default_branch = info['default_branch']
    conn.save(update_fields=['repo_full_name', 'default_branch'])
    return JsonResponse({'status': 'ok', 'repo': conn.repo_full_name, 'branch': conn.default_branch})


def github_disconnect(request):
    if not _github_guard(request):
        return _github_forbidden()
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)
    GitHubConnection.objects.filter(user=request.user).delete()
    return JsonResponse({'status': 'ok'})


def ai_github_send(request):
    if not _github_guard(request):
        return _github_forbidden()
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'detail': 'Invalid request method.'}, status=405)

    cache_key = f'ai_github_rate:{request.user.pk}'
    count = cache.get(cache_key, 0)
    if count >= AI_GITHUB_RATE_LIMIT:
        return JsonResponse({'status': 'error', 'detail': 'Too many GitHub requests — please wait a bit.'}, status=429)
    cache.set(cache_key, count + 1, AI_CHAT_RATE_WINDOW)

    conn = GitHubConnection.objects.filter(user=request.user).first()
    if not conn or not conn.repo_full_name:
        return JsonResponse({'status': 'error', 'detail': 'Connect a GitHub repo first.'}, status=400)

    payload = _parse_json_body(request)
    if not isinstance(payload, dict):
        return JsonResponse({'status': 'error', 'detail': 'Invalid request body.'}, status=400)
    message = str(payload.get('message', ''))[:AI_CHAT_MAX_MESSAGE_CHARS].strip()
    if not message:
        return JsonResponse({'status': 'error', 'detail': 'No instruction provided.'}, status=400)
    display_model_key = (
        ai_chat.CHATGPT_56_MODEL_KEY
        if payload.get('model') == ai_chat.CHATGPT_56_MODEL_KEY
        else 'github'
    )

    owner, _, repo = conn.repo_full_name.partition('/')
    branch = conn.default_branch or 'main'

    owner_filter = _ai_owner_filter(request)
    conversation_id = payload.get('conversation_id')
    conversation = AIConversation.objects.filter(owner_filter, pk=conversation_id).first() if conversation_id else None
    if conversation is None:
        conversation = AIConversation.objects.create(user=request.user, title=message[:AI_CONVERSATION_TITLE_CHARS])
    request.session[AI_CURRENT_CONVERSATION_SESSION_KEY] = conversation.id
    AIMessage.objects.create(conversation=conversation, role=AIMessage.ROLE_USER, content=message)
    conversation.updated_at = timezone.now()
    conversation.save(update_fields=['updated_at'])

    def finish(reply_text):
        AIMessage.objects.create(
            conversation=conversation, role=AIMessage.ROLE_ASSISTANT,
            content=reply_text, model_key=display_model_key,
        )
        return JsonResponse({
            'status': 'ok', 'reply': reply_text,
            'conversation_id': conversation.id, 'model_key': display_model_key,
        })

    try:
        file_paths = github_ops.get_tree(conn.access_token, owner, repo, branch)
    except github_ops.GitHubAPIError as e:
        return finish(f"Couldn't read the repository: {e}")

    wanted = ai_chat.github_select_files(message, file_paths)
    file_contents, file_shas = {}, {}
    for path in wanted:
        if github_ops.is_path_blocked(path):
            # The blocked list protects settings/migrations/secrets/etc from
            # writes below — it must cover reads too, or a blocked file's
            # contents (e.g. settings.py) would still get pasted into the
            # planning prompt even though it can never actually be changed.
            continue
        try:
            content, sha = github_ops.get_file(conn.access_token, owner, repo, path, branch)
        except github_ops.GitHubAPIError:
            continue
        if len(content) > AI_GITHUB_MAX_FILE_CHARS:
            content = content[:AI_GITHUB_MAX_FILE_CHARS] + '\n...[truncated]'
        file_contents[path] = content
        file_shas[path] = sha

    try:
        plan = ai_chat.github_plan_changes(message, file_paths, file_contents)
    except Exception as e:
        logger.exception("GitHub plan generation failed: %s", e)
        return finish("Something went wrong while planning the change — please try again or rephrase your request.")

    if not isinstance(plan, dict):
        return finish("The AI's response wasn't understood — please try rephrasing your request.")

    summary = str(plan.get('summary') or 'No changes were made.')
    commit_message = str(plan.get('commit_message') or message)[:200] or message[:200]
    operations = plan.get('operations') or []

    # Validate every proposed operation before touching the GitHub API at
    # all, so we know whether there's anything real to commit BEFORE
    # creating a working branch for it (never open an empty branch/PR).
    valid_ops, skipped = [], []
    for op in operations[:20]:
        if not isinstance(op, dict):
            continue
        action = op.get('action')
        path = str(op.get('path', '')).strip()
        if not path or github_ops.is_path_blocked(path):
            skipped.append(f"{path or '(blank path)'} — blocked")
            continue
        if action in ('update', 'create'):
            content = op.get('content')
            if not isinstance(content, str):
                skipped.append(f"{path} — no content given")
                continue
            valid_ops.append({'action': action, 'path': path, 'content': content})
        elif action == 'delete':
            valid_ops.append({'action': 'delete', 'path': path})
        else:
            skipped.append(f"{path} — unknown action '{action}'")

    applied, pr_url, work_branch = [], None, None
    if valid_ops:
        # Every change lands on a fresh branch + PR rather than a direct
        # commit to the default branch — an AI-proposed change (from a
        # small, non-specialized model, on a plain-English instruction) has
        # no business landing on main unreviewed. If nothing ends up
        # actually applied below, the branch is deleted again so this
        # doesn't litter the repo with empty branches.
        work_branch = f"ai/{timezone.now():%Y%m%d-%H%M%S}-{secrets.token_hex(3)}"
        try:
            base_sha = github_ops.get_branch_sha(conn.access_token, owner, repo, branch)
            github_ops.create_branch(conn.access_token, owner, repo, work_branch, base_sha)
        except github_ops.GitHubAPIError as e:
            return finish(f"Couldn't create a working branch for this change: {e}")

        for op in valid_ops:
            path = op['path']
            try:
                if op['action'] in ('update', 'create'):
                    sha = file_shas.get(path)
                    if sha is None and op['action'] == 'update':
                        try:
                            _, sha = github_ops.get_file(conn.access_token, owner, repo, path, work_branch)
                        except github_ops.GitHubAPIError:
                            sha = None
                    github_ops.upsert_file(
                        conn.access_token, owner, repo, path, op['content'], commit_message, work_branch, sha=sha,
                    )
                    applied.append(path)
                else:
                    sha = file_shas.get(path)
                    if sha is None:
                        try:
                            _, sha = github_ops.get_file(conn.access_token, owner, repo, path, work_branch)
                        except github_ops.GitHubAPIError:
                            sha = None
                    if sha is None:
                        skipped.append(f"{path} — not found")
                        continue
                    github_ops.delete_file(conn.access_token, owner, repo, path, commit_message, work_branch, sha)
                    applied.append(path)
            except github_ops.GitHubAPIError as e:
                skipped.append(f"{path} — {e}")

        if applied:
            try:
                pr = github_ops.create_pull_request(
                    conn.access_token, owner, repo, title=commit_message,
                    head=work_branch, base=branch, body=summary,
                )
                pr_url = pr.get('html_url')
            except github_ops.GitHubAPIError as e:
                skipped.append(f"(pull request could not be opened automatically: {e})")
        else:
            github_ops.delete_branch(conn.access_token, owner, repo, work_branch)

    lines = [summary]
    if applied:
        lines.append(f"\n**Proposed on `{work_branch}` — not yet merged:**")
        lines.extend(f"- {p}" for p in applied)
        if pr_url:
            lines.append(f"\nReview and merge: {pr_url}")
        else:
            lines.append(
                f"\nPushed to `{work_branch}` in {conn.repo_full_name}, but opening a pull "
                "request automatically failed — you can open one manually from that branch."
            )
    if skipped:
        lines.append("\n**Skipped:**")
        lines.extend(f"- {s}" for s in skipped)
    if not applied and not skipped and not plan.get('summary'):
        lines.append("\nNo file changes were made.")
    return finish('\n'.join(lines))
