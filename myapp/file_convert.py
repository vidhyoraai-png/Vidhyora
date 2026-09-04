"""Convert an uploaded file between document, spreadsheet and image formats.

Everything here runs locally with libraries already in requirements.txt — no
conversion service, no API key, no network call — so a conversion either works
offline or fails cleanly with a message the user can act on.

Scope is deliberately honest: these are **content** conversions, not
pixel-perfect layout ones. Turning a DOCX into a PDF without LibreOffice means
carrying the text, headings and lists across, not reproducing the original's
exact typography. The UI says so, because a conversion that silently drops a
letterhead is worse than one that told you it would.

``convert()`` is the single entry point; ``targets_for()`` says what a given
file can become, so the UI never offers a conversion that would then fail.
"""

import csv
import io
import os
import re

MAX_PDF_PAGES = 100
# Matches doc_extract.MAX_CSV_ROWS in spirit — a spreadsheet big enough to
# exceed this is a database export, not something to hand back through a chat
# attachment, and materialising it would blow the request's memory budget.
MAX_SHEET_ROWS = 20_000


class ConvertError(Exception):
    """A conversion that cannot be completed, with a user-facing reason."""


# Source extension -> what it can be turned into. Kept explicit rather than
# computed so the UI's dropdown and the server's validation can never disagree.
_CONVERSIONS = {
    'pdf': ('docx', 'txt'),
    'docx': ('pdf', 'txt'),
    'txt': ('pdf', 'docx'),
    'md': ('pdf', 'docx'),
    'csv': ('xlsx', 'pdf'),
    'xlsx': ('csv', 'pdf'),
    'png': ('pdf', 'jpg', 'webp'),
    'jpg': ('pdf', 'png', 'webp'),
    'jpeg': ('pdf', 'png', 'webp'),
    'webp': ('pdf', 'png', 'jpg'),
    'bmp': ('pdf', 'png', 'jpg', 'webp'),
}

_MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'csv': 'text/csv; charset=utf-8',
    'txt': 'text/plain; charset=utf-8',
    'png': 'image/png',
    'jpg': 'image/jpeg',
    'webp': 'image/webp',
}


def _extension(filename):
    return (os.path.splitext(filename or '')[1] or '').lstrip('.').lower()


def targets_for(filename):
    """Formats ``filename`` can be converted into (possibly empty)."""
    return list(_CONVERSIONS.get(_extension(filename), ()))


def mime_for(extension):
    return _MIME_TYPES.get((extension or '').lower(), 'application/octet-stream')


# ── source readers: every supported input becomes either text or table rows ──

def _pdf_to_text(data):
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ConvertError("That PDF couldn't be read — it may be corrupt.") from exc
    if getattr(reader, 'is_encrypted', False):
        raise ConvertError('That PDF is password-protected — remove the password and try again.')

    pages = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        try:
            pages.append(page.extract_text() or '')
        except Exception:
            continue
    text = '\n\n'.join(part for part in pages if part.strip())
    if not text.strip():
        # A scanned PDF is images with no text layer. Say that plainly rather
        # than handing back an empty document that looks like a silent success.
        raise ConvertError(
            "That PDF has no selectable text — it looks like a scan. "
            'Convert it to an image, or upload it in chat so it can be read with OCR.'
        )
    return text


def _docx_to_text(data):
    from docx import Document as DocxDocument

    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise ConvertError("That Word file couldn't be read — it may be corrupt.") from exc

    lines = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or '').strip()
        if not text:
            lines.append('')
            continue
        # Carry Word's heading levels across as Markdown so the PDF/text
        # writers below reproduce the document's structure instead of
        # flattening everything into identical paragraphs.
        style = (getattr(paragraph.style, 'name', '') or '').lower()
        heading = re.match(r'heading (\d)', style)
        if heading:
            lines.append('#' * min(int(heading.group(1)), 6) + ' ' + text)
        elif style.startswith('list'):
            lines.append('- ' + text)
        else:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or '').strip() for cell in row.cells]
            if any(cells):
                lines.append(' | '.join(cells))
    return '\n'.join(lines).strip()


def _xlsx_to_rows(data):
    from openpyxl import load_workbook

    try:
        workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except Exception as exc:
        raise ConvertError("That spreadsheet couldn't be read — it may be corrupt.") from exc
    sheet = workbook.active
    rows = []
    for row in sheet.iter_rows(values_only=True):
        rows.append(['' if value is None else str(value) for value in row])
        if len(rows) >= MAX_SHEET_ROWS:
            break
    workbook.close()
    return rows


def _csv_to_rows(data):
    text = _decode_text(data)
    return [row for row in csv.reader(io.StringIO(text))][:MAX_SHEET_ROWS]


def _decode_text(data):
    for encoding in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


# ── writers ──

_PDF_PAGE_SIZE = (612, 792)      # US Letter, in points
_PDF_MARGIN = 56
_PDF_BODY_SIZE = 11
_PDF_HEADING_SIZES = {1: 20, 2: 17, 3: 15, 4: 13, 5: 12, 6: 11}
_PDF_LINE_SPACING = 1.35
_PDF_FONT = 'helv'
_PDF_FONT_BOLD = 'hebo'


def _wrap_line(fitz_module, text, fontsize, max_width, fontname):
    """Greedy word-wrap using the font's real glyph widths."""
    words = text.split(' ')
    lines = []
    current = ''
    for word in words:
        candidate = f'{current} {word}'.strip()
        if not current or fitz_module.get_text_length(
            candidate, fontname=fontname, fontsize=fontsize,
        ) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def text_to_pdf_bytes(content):
    """Render Markdown-ish text (headings, bullets, numbered lists) as a PDF.

    Shared by the AI's own file generation (views._ai_pdf_bytes) and by the
    conversions here, so both produce identically-formatted documents.

    PyMuPDF is imported lazily, the way doc_extract.py does it — a missing
    optional dependency must degrade to a clear message, never an import-time
    crash that takes the whole site down.
    """
    try:
        import fitz
    except ImportError as exc:
        raise ConvertError('PDF generation is temporarily unavailable.') from exc

    width, height = _PDF_PAGE_SIZE
    max_width = width - 2 * _PDF_MARGIN
    doc = fitz.open()
    page = doc.new_page(width=width, height=height)
    y = _PDF_MARGIN

    def emit(text, fontsize, fontname, indent=0, gap_after=6):
        nonlocal page, y
        for line in _wrap_line(fitz, text, fontsize, max_width - indent, fontname):
            if y + fontsize > height - _PDF_MARGIN:
                page = doc.new_page(width=width, height=height)
                y = _PDF_MARGIN
            page.insert_text(
                (_PDF_MARGIN + indent, y + fontsize), line,
                fontsize=fontsize, fontname=fontname,
            )
            y += fontsize * _PDF_LINE_SPACING
        y += gap_after

    for raw_line in (content or '').splitlines():
        line = raw_line.strip()
        if not line:
            y += _PDF_BODY_SIZE * 0.5
            continue
        heading = re.match(r'^(#{1,6})\s+(.+)$', line)
        bullet = re.match(r'^[-*]\s+(.+)$', line)
        numbered = re.match(r'^(\d+[.)])\s+(.+)$', line)
        if heading:
            emit(heading.group(2), _PDF_HEADING_SIZES[min(len(heading.group(1)), 6)], _PDF_FONT_BOLD)
        elif bullet:
            # A plain hyphen, not a Unicode bullet: the base-14 PDF fonts here
            # don't reliably round-trip "•" through every reader/extractor.
            emit(f'- {bullet.group(1)}', _PDF_BODY_SIZE, _PDF_FONT, indent=14)
        elif numbered:
            emit(f'{numbered.group(1)} {numbered.group(2)}', _PDF_BODY_SIZE, _PDF_FONT, indent=14)
        else:
            emit(line, _PDF_BODY_SIZE, _PDF_FONT)

    data = doc.tobytes()
    doc.close()
    return data


def text_to_docx_bytes(content):
    """Render Markdown-ish text as a real DOCX package."""
    from docx import Document as WordDocument

    document = WordDocument()
    for raw_line in (content or '').splitlines():
        line = raw_line.strip()
        heading = re.match(r'^(#{1,6})\s+(.+)$', line)
        bullet = re.match(r'^[-*]\s+(.+)$', line)
        numbered = re.match(r'^\d+[.)]\s+(.+)$', line)
        if heading:
            document.add_heading(heading.group(2), level=min(len(heading.group(1)), 9))
        elif bullet:
            document.add_paragraph(bullet.group(1), style='List Bullet')
        elif numbered:
            document.add_paragraph(numbered.group(1), style='List Number')
        else:
            document.add_paragraph(raw_line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def rows_to_xlsx_bytes(rows):
    """Write table rows as a real XLSX workbook, numbers stored as numbers."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise ConvertError('Excel generation is temporarily unavailable.') from exc

    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append([_cell_value(value) for value in row])
    if rows:
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for index, column in enumerate(sheet.columns, start=1):
            longest = max((len(str(cell.value or '')) for cell in column), default=0)
            sheet.column_dimensions[get_column_letter(index)].width = min(max(longest + 2, 8), 60)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _cell_value(value):
    """Numeric-looking cells become real numbers; everything else stays text."""
    text = str(value or '').strip()
    if not text:
        return ''
    # Narrow on purpose: a currency symbol, percent sign or thousands separator
    # means reinterpreting would change what the user actually wrote.
    if re.fullmatch(r'-?\d+', text):
        return int(text)
    if re.fullmatch(r'-?\d*\.\d+', text):
        return float(text)
    return text


def rows_to_csv_bytes(rows):
    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator='\n')
    writer.writerows(rows)
    return buffer.getvalue().encode('utf-8')


def _rows_to_text(rows):
    return '\n'.join(' | '.join(row) for row in rows)


def _image_to_pdf_bytes(data):
    """One image, scaled to fit a single page with the aspect ratio intact."""
    from PIL import Image

    try:
        image = Image.open(io.BytesIO(data))
        image.load()
    except Exception as exc:
        raise ConvertError("That image couldn't be read — it may be corrupt.") from exc
    # PDF has no alpha channel; flatten onto white so a transparent PNG
    # doesn't come out with a black background.
    if image.mode in ('RGBA', 'LA', 'P'):
        image = image.convert('RGBA')
        canvas = Image.new('RGB', image.size, (255, 255, 255))
        canvas.paste(image, mask=image.split()[-1])
        image = canvas
    else:
        image = image.convert('RGB')

    buffer = io.BytesIO()
    image.save(buffer, format='PDF', resolution=100.0)
    return buffer.getvalue()


def _image_to_image_bytes(data, target):
    from PIL import Image

    try:
        image = Image.open(io.BytesIO(data))
        image.load()
    except Exception as exc:
        raise ConvertError("That image couldn't be read — it may be corrupt.") from exc

    pillow_format = {'jpg': 'JPEG', 'png': 'PNG', 'webp': 'WEBP'}[target]
    if pillow_format == 'JPEG' and image.mode in ('RGBA', 'LA', 'P'):
        image = image.convert('RGBA')
        canvas = Image.new('RGB', image.size, (255, 255, 255))
        canvas.paste(image, mask=image.split()[-1])
        image = canvas
    elif pillow_format == 'JPEG':
        image = image.convert('RGB')

    buffer = io.BytesIO()
    image.save(buffer, format=pillow_format)
    return buffer.getvalue()


def convert(data, source_name, target):
    """Convert ``data`` (bytes of ``source_name``) into ``target``.

    Returns ``(bytes, filename, mimetype)``. Raises ConvertError with a
    user-facing message for anything unsupported or unreadable.
    """
    source = _extension(source_name)
    target = (target or '').lower().lstrip('.')
    if target == 'jpeg':
        target = 'jpg'
    if not data:
        raise ConvertError('That file is empty.')
    if source not in _CONVERSIONS:
        raise ConvertError(f'{source.upper() or "That file type"} files cannot be converted here.')
    if target not in _CONVERSIONS[source]:
        allowed = ', '.join(_CONVERSIONS[source]).upper()
        raise ConvertError(f'A {source.upper()} can only be converted to: {allowed}.')

    if source in ('png', 'jpg', 'jpeg', 'webp', 'bmp'):
        payload = _image_to_pdf_bytes(data) if target == 'pdf' else _image_to_image_bytes(data, target)
    elif source in ('csv', 'xlsx'):
        rows = _csv_to_rows(data) if source == 'csv' else _xlsx_to_rows(data)
        if target == 'xlsx':
            payload = rows_to_xlsx_bytes(rows)
        elif target == 'csv':
            payload = rows_to_csv_bytes(rows)
        else:
            payload = text_to_pdf_bytes(_rows_to_text(rows))
    else:
        if source == 'pdf':
            text = _pdf_to_text(data)
        elif source == 'docx':
            text = _docx_to_text(data)
        else:
            text = _decode_text(data)
        if target == 'pdf':
            payload = text_to_pdf_bytes(text)
        elif target == 'docx':
            payload = text_to_docx_bytes(text)
        else:
            payload = text.encode('utf-8')

    stem = os.path.splitext(os.path.basename(source_name or 'converted'))[0] or 'converted'
    stem = re.sub(r'[^A-Za-z0-9_.-]+', '-', stem).strip('.-')[:80] or 'converted'
    return payload, f'{stem}.{target}', mime_for(target)
